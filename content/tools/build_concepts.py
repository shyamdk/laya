#!/usr/bin/env python3
"""Two beginner-friendly concept guides — one per chapter — covering every idea
used in the question bank. Methods follow NCERT Ganita Prakash Grade 8."""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from build_doc import rich, NAVY, GREY, GREEN

DIR = "/Users/shyamdk/Developer/personal/laya/data/question-papers/"
AMBER = RGBColor(0xB2, 0x5E, 0x00)
BLUE = RGBColor(0x0D, 0x47, 0xA1)

# section number -> (tier, why).  Tiers come from topics.py — counted from the real papers.
TIER1 = {
 2: ("BASICS", "You cannot do anything else without these."),
 3: ("**", "Identifying / ruling out perfect squares came up in 4 exams."),
 4: ("**", "Odd-number patterns came up in 4 exams."),
 5: ("**", "'How many numbers between two squares' came up in 4 exams."),
 6: ("BASICS", "The idea behind every square-root question."),
 7: ("***", "Prime factorisation is the engine behind decimals, LCM and perfect-square questions."),
 8: ("**", "Estimation rescues you when factorising is ugly."),
 9: ("***", "THE most asked topic — 10 questions across 7 of the 11 exams."),
 10: ("**", "Smallest multiplier/divisor came up in 4 exams."),
 11: ("*", "LCM-type square/cube questions came up in 3 exams."),
 12: ("**", "'Rows = columns' word problems came up in 4 exams."),
 13: ("BASICS", "Needed for every cube question."),
 14: ("***", "This table is what makes cube-root guessing work — 7 exams."),
 15: ("*", "Cube root by prime factorisation came up in 2 exams."),
 16: ("***", "Cube root by estimation — 7 questions across 7 of the 11 exams."),
 17: ("*", "Making a perfect cube came up in 3 exams."),
 18: ("*", "Ratio + sum of cubes came up in 2 exams."),
 19: ("**", "Not in your book — yet it still appeared in 4 exams. Worth knowing."),
}
TIER2 = {
 2: ("BASICS", "Recognising 32 = 2^5 or 81 = 3^4 unlocks most questions."),
 3: ("***", "THE most asked topic — 13 questions across 8 of the 9 exams."),
 4: ("*", "Asked directly twice, but used inside many other questions."),
 5: ("**", "Negative exponents came up in 5 exams."),
 6: ("**", "The 'flip the fraction' move appears throughout."),
 7: ("***", "'Find x' came up in 8 questions across 7 of the 9 exams."),
 8: ("*", "Converting to standard form came up in 3 exams."),
 9: ("**", "Standard-form arithmetic came up in 6 exams."),
 10: ("**", "Standard-form arithmetic came up in 6 exams."),
 11: ("", "Beyond your book, and appeared only once."),
}
TIER_COL = {"***": RGBColor(0xC6, 0x28, 0x28), "**": AMBER, "*": GREY, "BASICS": BLUE, "": GREY}
TIER_LBL = {"***": "*** MUST KNOW", "**": "** IMPORTANT", "*": "* SEEN SOMETIMES",
            "BASICS": "FOUNDATION", "": ""}

TEAL = RGBColor(0x00, 0x69, 0x5C)

# section number -> a real thing a Grade 8 student actually meets
REAL1 = {
 1: "A Rubik's cube is 3 x 3 x 3 = 3^{3} = 27 small cubies. A carrom board is 8 x 8 = 8^{2} = 64 squares. You have been using squares and cubes since you could count.",
 2: "Chess and carrom boards are 8 x 8 = 64. A Ludo board's home column, a rangoli grid, floor tiles in your classroom - all squares. If you can spot 64, 144 or 225 instantly, you are already halfway through most questions.",
 3: "Your friend says the area of a square badminton court is 8123 sq ft. You can tell her she is wrong in two seconds: a square never ends in 3. This is the fastest 'check my answer' tool you own - use it on your own answers before you hand the paper in.",
 4: "Think of building a square rangoli. To grow a 3 x 3 rangoli into a 4 x 4, you add an L-shaped border of 7 dots. 9 + 7 = 16. Each new border is the next odd number - that is why squares are made of odd numbers.",
 5: "Between 12^{2} = 144 and 13^{2} = 169 there is no perfect square at all - just 24 ordinary numbers. So if a shopkeeper claims she has 150 tiles and can lay them in a perfect square, she cannot.",
 6: "Your bedroom floor is square and its area is 144 sq ft. How long is one wall? That is exactly a square root: 12 ft. Area -> side is the single most common real use of square roots.",
 7: "A square sheet of paper has an area of 7056 sq cm. To cut a border you need its side length - so you need √7056 = 84 cm. Prime factorisation gets you there with no calculator.",
 8: "Your teacher asks for √1936 and there is no calculator in the exam hall. Trapping it between 40^{2} and 50^{2}, then using the last digit, gives you 44 in about fifteen seconds.",
 9: "A square plot of land measures 125.44 sq m. The fencing contractor needs the side length: √125.44 = 11.2 m. Decimals turn up the moment you leave the textbook and touch real measurements - which is exactly why this is the most-asked topic.",
 10: "You have 2352 square tiles and want to lay them in a perfect square patch with none left over. You cannot - but buy 3 more sets and you get 7056 = 84 x 84. That is what 'smallest number to multiply by' really means.",
 11: "You want the smallest square hall that can be split evenly into rooms of 10, 20 and 28 units. Answer: 4900 sq units, i.e. 70 x 70.",
 12: "This is the school assembly, the Republic Day parade, a dance formation, or seats in an auditorium - anything where rows must equal columns. 2016 dancers cannot make a square; add 9 more and you get a perfect 45 x 45.",
 13: "The volume of a cubical water tank, a box of sugar cubes, a dice, a Rubik's cube. If a cube-shaped tank holds 9261 litres, each side is ∛9261 = 21 units.",
 14: "Use this as a party trick. Ask someone to cube any two-digit number on a calculator and read out the answer. Because every digit cubes to a different ending, you can name their number back to them in three seconds.",
 15: "A cubical shipping carton has a volume of 17576 cubic cm. What is its height? ∛17576 = 26 cm - so it fits on a 30 cm shelf.",
 16: "Same party trick as above, done properly: 658503 - split it 658 | 503, the 3 tells you it ends in 7, and 658 sits between 8^{3} and 9^{3} so it starts with 8. Answer 87, faster than they can type it.",
 17: "You have boxes of 15 x 30 x 15 cm. How many do you need to stack into a perfect cube? Four - giving a 30 cm cube. Warehouse and packing problems are exactly this.",
 18: "Three cousins share Diwali money in the ratio 2 : 3 : 4. Once you know the sum of the cubes, one line of algebra (99x^{3}) gives everyone's share.",
 19: "A ladder leaning on a wall, the diagonal of a TV screen, or the shortest path across a rectangular park. A 'a 32-inch TV' means its diagonal is 32 inches - and its height and width form a triplet with it.",
}
REAL2 = {
 1: "You forward a message to 2 friends, they each forward to 2 more, and so on. After 30 rounds that is 2^{30} - over a BILLION people. This is exactly why a video 'goes viral', and why a rumour in school spreads by lunch break.",
 2: "Computer memory is built on powers of 2. 2^{10} = 1024 = 1 KB. 1 MB is about a thousand KB, 1 GB is about a thousand MB. When your phone says '64 GB', that 64 is 2^{6} - the whole device is built out of this table.",
 3: "Doubling your pocket money for 3 days, then doubling it for 2 more days, is the same as doubling it for 5 days: 2^{3} x 2^{2} = 2^{5}. That is Law 1, and you already knew it.",
 4: "Anything to the power zero is 1 because you have not multiplied at all - you still have your starting single unit. Zero doublings of Rs 100 leaves you with Rs 100, not Rs 0.",
 5: "Negative powers mean HALVING instead of doubling. 2^{-1} is half, 2^{-2} is a quarter. Zooming OUT on a map, cutting a cake in half again and again, or a medicine dose halving in your blood every few hours.",
 6: "A recipe for 4 people scaled DOWN. Flipping the fraction is what your brain already does when you say 'two-thirds as much' versus 'one-and-a-half times as much'.",
 7: "This is just detective work: you know the answer is 81 and the base is 3, so how many 3s were multiplied? Bacteria doubling, money compounding, a rumour spreading - 'how many rounds did it take?' is always this question.",
 8: "India's population is about 1.4 x 10^{9} people. A coronavirus is about 1 x 10^{-7} m across. Your phone camera photo is 4 x 10^{6} pixels. Nobody writes those zeros out - scientists and engineers use standard form all day.",
 9: "The distance to the Moon is 3.84 x 10^{8} m and to the Sun is 1.5 x 10^{11} m. To find how much FURTHER the Sun is, you must first line the powers up - just like you cannot add 1/2 and 1/3 until the denominators match.",
 10: "'The Earth is 88 times heavier than the Moon.' 'The Sun is 400 times further than the Moon.' Every time a science documentary says 'X times bigger', somebody divided two numbers in standard form.",
}

SQ = [(f"{n}", f"{n*n}") for n in range(1, 31)]
CU = [(f"{n}", f"{n**3}") for n in range(1, 21)]

# =========================== CHAPTER 1 ===========================
C1 = [
("Start here — what this chapter is really about", [
 ("p", "Two ideas, and everything else grows out of them."),
 ("p", "SQUARE a number = multiply it by itself.  5 squared = 5 × 5 = 25.  We write it 5^{2}."),
 ("p", "CUBE a number = multiply it by itself three times.  5 cubed = 5 × 5 × 5 = 125.  We write it 5^{3}."),
 ("p", "The names come from shapes. A square tile of side 5 has area 5 × 5 = 25. A cube box of side 5 holds 5 × 5 × 5 = 125 little cubes. That is all the words 'square' and 'cube' mean here."),
 ("p", "Most of this chapter is the REVERSE journey: somebody hands you 25 and you must work back to 5. That backwards step is called finding the square root. Same for cubes — from 125 back to 5 is the cube root."),
 ("key", "Square → forwards.  Square root → backwards.  That is the whole game."),
]),
("The squares you should simply know", [
 ("p", "Learn these. Nearly every question becomes easy once you can recognise these numbers on sight."),
 ("tbl", (["n", "n^{2}", "n", "n^{2}", "n", "n^{2}"],
          [[SQ[i][0], SQ[i][1], SQ[i+10][0], SQ[i+10][1], SQ[i+20][0], SQ[i+20][1]] for i in range(10)])),
 ("p", "A number that appears in the n^{2} column — 1, 4, 9, 16, 25, … — is called a PERFECT SQUARE."),
]),
("Trick 1 — spotting a number that CANNOT be a square", [
 ("p", "Look at the last digit of every square: 1, 4, 9, 16, 25, 36, 49, 64, 81, 100…"),
 ("p", "They end only in 0, 1, 4, 5, 6 or 9. A square NEVER ends in 2, 3, 7 or 8."),
 ("eg", ["Is 1253687 a perfect square?",
         "It ends in 7. Squares never end in 7.",
         "So no — and you knew it in two seconds, without any calculation."]),
 ("warn", "This test only works ONE WAY. If a number ends in 2, 3, 7 or 8 it is definitely not a square. But ending in 6 does NOT make it a square — 26 ends in 6 and is not a square. The test can rule numbers OUT, never rule them IN."),
 ("p", "Zeros follow a rule too: a square always ends in an EVEN number of zeros. So 6250 (one zero) cannot be a square, but 4900 (two zeros) can be — it is 70^{2}."),
]),
("Trick 2 — squares are made of odd numbers", [
 ("p", "Add up odd numbers starting from 1 and watch what happens:"),
 ("eg", ["1 = 1 = 1^{2}",
         "1 + 3 = 4 = 2^{2}",
         "1 + 3 + 5 = 9 = 3^{2}",
         "1 + 3 + 5 + 7 = 16 = 4^{2}"]),
 ("key", "The sum of the first n odd numbers is exactly n^{2}."),
 ("p", "This gives a free method: to test if a number is a square, keep SUBTRACTING 1, 3, 5, 7, … If you land exactly on 0, it is a square, and the number of subtractions is the root."),
 ("eg", ["Is 25 a square?",
         "25 − 1 = 24,  24 − 3 = 21,  21 − 5 = 16,  16 − 7 = 9,  9 − 9 = 0",
         "We landed on 0 after 5 subtractions, so 25 is a square and √25 = 5."]),
]),
("How many numbers sit between two squares?", [
 ("p", "Between 3^{2} = 9 and 4^{2} = 16 the numbers are 10, 11, 12, 13, 14, 15 — that is 6 numbers, and 6 = 2 × 3."),
 ("key", "Between n^{2} and (n+1)^{2} there are exactly 2n numbers.  (Or: between two squares p and q, there are q − p − 1 numbers.)"),
 ("eg", ["How many numbers lie between 18^{2} and 19^{2}?",
         "2 × 18 = 36 numbers."]),
]),
("Square root — the backwards step", [
 ("p", "√25 = 5 means: what number, multiplied by itself, gives 25?"),
 ("p", "Strictly, both 5 and −5 work, because (−5) × (−5) = 25 too. Your book says: in this chapter we only use the POSITIVE one. So √25 = 5."),
 ("p", "There are three ways to find a square root. Use whichever suits the number."),
]),
("Method A — square root by prime factorisation (the workhorse)", [
 ("p", "Break the number into prime factors, then split them into TWO IDENTICAL GROUPS. One group is the answer."),
 ("steps", ["1. Divide by 2 as long as you can, then 3, then 5, then 7 … until only primes remain.",
            "2. Pair up identical primes.",
            "3. Take ONE from each pair and multiply them together."]),
 ("eg", ["Find √7056.",
         "7056 = 2 × 3528 = 2 × 2 × 1764 = 2 × 2 × 2 × 882 = 2 × 2 × 2 × 2 × 441",
         "and 441 = 3 × 147 = 3 × 3 × 49 = 3 × 3 × 7 × 7",
         "So 7056 = 2^{4} × 3^{2} × 7^{2}",
         "Pairs: (2×2) (2×2) (3×3) (7×7)",
         "Take one from each pair: 2 × 2 × 3 × 7 = 84",
         "√7056 = 84 ✔"]),
 ("key", "If EVERY prime can be paired up, the number is a perfect square. If one prime is left lonely, it is not."),
]),
("Method B — square root by estimation (when factorising is painful)", [
 ("p", "This is the method your book uses for √1936 (page 9). Two ideas, used together."),
 ("steps", ["1. TRAP it. Find the two whole squares it sits between.",
            "2. Use the LAST DIGIT to narrow the choice down.",
            "3. Test the one or two candidates by squaring them."]),
 ("eg", ["Find √1936.",
         "Trap: 40^{2} = 1600 and 50^{2} = 2500, so the root is between 40 and 50.",
         "Last digit: 1936 ends in 6. Only 4^{2} = 16 and 6^{2} = 36 end in 6, so the root ends in 4 or 6.",
         "Candidates: 44 or 46. Test 44: 44 × 44 = 1936 ✔",
         "√1936 = 44"]),
 ("eg", ["Find √20 correct to two decimal places (20 is NOT a perfect square).",
         "4^{2} = 16, 5^{2} = 25 → between 4 and 5.",
         "4.4^{2} = 19.36 (too small), 4.5^{2} = 20.25 (too big) → between 4.4 and 4.5",
         "4.47^{2} = 19.9809 (just under), 4.48^{2} = 20.0704 (just over)",
         "So √20 ≈ 4.47"]),
]),
("Square roots of decimals — clear the decimal first", [
 ("p", "Do not panic at the decimal point. Turn it into a whole number, take the root, then put the point back."),
 ("steps", ["1. Write the decimal as a fraction over 100 (or 10000, …).",
            "2. Take the square root of the top and the bottom separately.",
            "3. √100 = 10, √10000 = 100 — easy."]),
 ("eg", ["Find √125.44",
         "125.44 = 12544 / 100",
         "12544 = 2^{8} × 7^{2}, so √12544 = 2^{4} × 7 = 112",
         "√125.44 = 112 / 10 = 11.2",
         "Always check: 11.2 × 11.2 = 125.44 ✔"]),
 ("p", "Useful shortcut to remember: squaring a decimal DOUBLES the number of decimal places, so a square root HALVES them. √0.0625 = 0.25 (4 places → 2 places)."),
]),
("Making a number into a perfect square", [
 ("p", "A very common exam question: 'What is the smallest number to multiply / divide by, to make this a perfect square?'"),
 ("steps", ["1. Prime factorise.",
            "2. Look for the LONELY primes — the ones with no partner.",
            "3. To MULTIPLY: multiply by exactly those lonely primes (giving each a partner).",
            "4. To DIVIDE: divide by exactly those lonely primes (throwing the lonely ones away)."]),
 ("eg", ["Smallest number to multiply 2352 by, to get a perfect square?",
         "2352 = 2^{4} × 3 × 7^{2}",
         "The 2s are paired ✔, the 7s are paired ✔, but 3 is lonely.",
         "Multiply by 3:  2352 × 3 = 7056 = 84^{2} ✔"]),
 ("eg", ["Smallest number to divide 486 by, to get a perfect square?",
         "486 = 2 × 3^{5}. The 2 is lonely, and 3^{5} has one lonely 3 (3^{4} pairs up, one 3 left over).",
         "Throw away 2 × 3 = 6:  486 ÷ 6 = 81 = 9^{2} ✔"]),
]),
("Smallest square divisible by several numbers", [
 ("steps", ["1. Find the LCM of the given numbers.",
            "2. Prime factorise the LCM.",
            "3. Multiply by whatever lonely primes remain, to pair everything up."]),
 ("eg", ["Smallest square number divisible by 10, 20 and 28.",
         "LCM(10, 20, 28) = 140 = 2^{2} × 5 × 7",
         "The 2s pair up ✔. But 5 and 7 are both lonely.",
         "Multiply by 5 × 7 = 35:  140 × 35 = 4900 = 70^{2} ✔"]),
]),
("The 'rows = columns' word problems", [
 ("p", "Soldiers in a square parade, chairs in a square hall, flowers in a square bed — the total must be a PERFECT SQUARE. Only one question matters: which squares does the given number sit between?"),
 ("warn", "Read the question very carefully. There are two opposite versions and it is the commonest place to lose marks."),
 ("eg", ["ADDING (how many MORE do we need?) — go UP to the next square.",
         "9450 people. 97^{2} = 9409 (too few), 98^{2} = 9604.",
         "We must reach 9604, so we need 9604 − 9450 = 154 more people, with 98 in each row."]),
 ("eg", ["LEFT OVER (how many are UNUSED?) — stay DOWN at the square below.",
         "13695 chairs. 117^{2} = 13689, 118^{2} = 13924 (too many).",
         "The biggest square we can build is 13689, so 13695 − 13689 = 6 chairs are left over."]),
]),
("Cubes — the same story, one dimension up", [
 ("p", "n^{3} = n × n × n. Learn these; they are used constantly."),
 ("tbl", (["n", "n^{3}", "n", "n^{3}"],
          [[CU[i][0], CU[i][1], CU[i+10][0], CU[i+10][1]] for i in range(10)])),
 ("p", "Cubes also come from odd numbers, in blocks:  1 = 1^{3};  3 + 5 = 8 = 2^{3};  7 + 9 + 11 = 27 = 3^{3};  21 + 23 + 25 + 27 + 29 = 125 = 5^{3}."),
 ("key", "The cube of n is the sum of n consecutive odd numbers."),
]),
("The cube last-digit table — a small piece of magic", [
 ("p", "For squares, the last digit was a weak clue. For CUBES it is a perfect fingerprint — every digit gives a different ending:"),
 ("tbl", (["last digit of n", "0","1","2","3","4","5","6","7","8","9"],
          [["last digit of n^{3}", "0","1","8","7","4","5","6","3","2","9"]])),
 ("p", "Read it backwards and you can name the last digit of any cube root instantly. A cube ending in 3 must have a root ending in 7. A cube ending in 8 must have a root ending in 2."),
]),
("Cube root — by prime factorisation", [
 ("p", "Same as squares, but group the primes in THREES instead of pairs."),
 ("eg", ["Find ∛17576.",
         "17576 = 2 × 8788 = 2 × 2 × 4394 = 2 × 2 × 2 × 2197",
         "and 2197 = 13 × 13 × 13",
         "So 17576 = 2^{3} × 13^{3}",
         "Groups of three: (2 × 2 × 2) and (13 × 13 × 13)",
         "Take one from each group: 2 × 13 = 26",
         "∛17576 = 26 ✔"]),
]),
("Cube root — by guessing (fast, and very impressive)", [
 ("p", "For a 4-, 5- or 6-digit cube you can write the answer down without factorising at all."),
 ("steps", ["1. Split the number into groups of THREE digits, starting from the right.",
            "2. The LAST group gives the unit digit of the answer — use the magic table above.",
            "3. The FIRST group gives the tens digit — find which two cubes it sits between, and take the SMALLER one."]),
 ("eg", ["Find ∛658503.",
         "Split:  658 | 503",
         "Last group 503 ends in 3 → the root ends in 7 (because 7^{3} = 343 ends in 3).",
         "First group is 658. Now 8^{3} = 512 and 9^{3} = 729, and 512 < 658 < 729 → take the smaller, 8.",
         "So the root is 87.  Check: 87^{3} = 658503 ✔"]),
]),
("Making a number into a perfect cube", [
 ("p", "Exactly like the square version, but now every prime needs to appear in a group of THREE."),
 ("eg", ["Is 2560 a perfect cube? If not, what is the smallest multiplier?",
         "2560 = 2^{9} × 5",
         "2^{9} is fine (that is three groups of three). But 5 appears once — it needs two more.",
         "Multiply by 5^{2} = 25:  2560 × 25 = 64000 = 40^{3} ✔"]),
 ("eg", ["Find x if 1323/x is a perfect cube.",
         "1323 = 3^{3} × 7^{2}",
         "3^{3} is a perfect group ✔. The 7^{2} can never be completed by dividing, so it must go entirely.",
         "x = 7^{2} = 49, leaving 1323 ÷ 49 = 27 = 3^{3} ✔"]),
]),
("Ratio problems with cubes", [
 ("p", "'Three numbers are in the ratio 2 : 3 : 4 and the sum of their cubes is …' — these look scary and are actually the same two lines every time."),
 ("steps", ["1. Call the numbers 2x, 3x, 4x.",
            "2. Cube each: 8x^{3}, 27x^{3}, 64x^{3}. They always add to 99x^{3}.",
            "3. Set 99x^{3} equal to the given total, solve for x, then write out 2x, 3x, 4x."]),
 ("eg", ["Sum of cubes is 131769.",
         "99x^{3} = 131769  ⇒  x^{3} = 1331  ⇒  x = 11",
         "The numbers are 22, 33 and 44."]),
]),
("⚠ Beyond your textbook — Pythagorean triplets", [
 ("warn", "This is NOT in your Ganita Prakash chapter, but some old exam papers use it, so here it is."),
 ("p", "Three whole numbers form a Pythagorean triplet when a^{2} + b^{2} = c^{2} — the sides of a right-angled triangle."),
 ("p", "You can always solve these with just your squares table. If the biggest number is 17, then a^{2} + b^{2} = 289. Now hunt: 289 − 64 = 225, and 225 = 15^{2}. So the triplet is 8, 15, 17."),
 ("p", "Triplets worth recognising on sight: (3,4,5), (5,12,13), (7,24,25), (8,15,17), (9,40,41), (10,24,26), (12,35,37), (16,63,65)."),
]),
("Mistakes to avoid", [
 ("bullets", ["Ending in 6 does not make a number a square. The last-digit test can only rule numbers OUT.",
              "√ and ∛ are not the same as ÷2 and ÷3. √36 = 6, not 18.",
              "A cube root of a negative number is negative: ∛(−64) = −4. (A square root of a negative number does not exist here.)",
              "In 'rows = columns' problems, ADDING means go up to the next square; LEFT OVER means stay at the square below.",
              "When making a perfect square, multiply by the LONELY primes only — not by the whole number.",
              "Always check your root by squaring (or cubing) it back. It takes five seconds and catches almost every slip."]),
]),
("Quick self-test", [
 ("check", [("Is 4728 a perfect square? How do you know instantly?", "No — it ends in 8, and squares never end in 2, 3, 7 or 8."),
            ("√7056 = ?", "84  (7056 = 2^{4} × 3^{2} × 7^{2}, so take 2^{2} × 3 × 7)"),
            ("∛9261 = ?", "21  (9261 = 3^{3} × 7^{3})"),
            ("How many numbers lie between 12^{2} and 13^{2}?", "2 × 12 = 24"),
            ("Smallest number to multiply 8820 by, to get a perfect square?", "5.  8820 = 2^{2} × 3^{2} × 5 × 7^{2} — only the 5 is lonely. The root then is 210."),
            ("A cube ends in 2. What does its cube root end in?", "8  (because 8^{3} = 512 ends in 2)"),
            ("√54.76 = ?", "7.4  (5476 = 2^{2} × 37^{2}, so √5476 = 74, then divide by 10)")]),
]),
]

# =========================== CHAPTER 2 ===========================
C2 = [
("Start here — what an exponent actually is", [
 ("p", "An exponent is nothing but SHORTHAND for multiplying a number by itself again and again."),
 ("p", "Instead of writing 5 × 5 × 5 × 5, we write 5^{4}. That is all."),
 ("p", "In 5^{4}:  the 5 is the BASE (the number being multiplied), and the 4 is the EXPONENT or POWER (how many times). We read it '5 to the power 4'."),
 ("key", "n^{a} means n multiplied by itself a times."),
 ("p", "Why bother? Because things that keep doubling or tripling get enormous astonishingly fast. Fold a paper 46 times and — in theory — it reaches the Moon. Exponents are how we write such numbers without a page full of zeros."),
]),
("Powers you should recognise instantly", [
 ("tbl", (["", "^{1}", "^{2}", "^{3}", "^{4}", "^{5}", "^{6}"],
          [["2", "2", "4", "8", "16", "32", "64"],
           ["3", "3", "9", "27", "81", "243", "729"],
           ["5", "5", "25", "125", "625", "3125", "15625"],
           ["10", "10", "100", "1000", "10000", "100000", "1000000"]])),
 ("p", "Also handy: 2^{7} = 128, 2^{8} = 256, 2^{9} = 512, 2^{10} = 1024."),
 ("p", "Spotting that 32 is 2^{5}, or that 81 is 3^{4}, is what turns a hard-looking question into an easy one."),
]),
("The five laws — and why each one is obvious", [
 ("p", "Do not memorise these blindly. Each one is just counting."),
 ("p", "LAW 1 — same base, MULTIPLYING → ADD the exponents.   n^{a} × n^{b} = n^{a+b}"),
 ("eg", ["2^{3} × 2^{2} = (2×2×2) × (2×2) = 2×2×2×2×2 = 2^{5}",
         "Three 2s and two more 2s make five 2s. That is all 'add the exponents' means."]),
 ("p", "LAW 2 — same base, DIVIDING → SUBTRACT the exponents.   n^{a} ÷ n^{b} = n^{a−b}"),
 ("eg", ["2^{5} ÷ 2^{3} = (2×2×2×2×2) / (2×2×2) = 2×2 = 2^{2}",
         "Three of the 2s cancel out. Five take away three leaves two."]),
 ("p", "LAW 3 — a power OF a power → MULTIPLY the exponents.   (n^{a})^{b} = n^{a×b}"),
 ("eg", ["(2^{3})^{2} = 2^{3} × 2^{3} = 2^{6}",
         "Two lots of three 2s = six 2s."]),
 ("p", "LAW 4 — same exponent, different bases → combine the bases.   n^{a} × m^{a} = (n × m)^{a}   and   n^{a} ÷ m^{a} = (n ÷ m)^{a}"),
 ("eg", ["8^{4} × (3/8)^{4} = (8 × 3/8)^{4} = 3^{4} = 81",
         "Because the exponents match, you may combine the bases first — which made this collapse to nothing."]),
 ("warn", "Laws 1, 2 and 3 need the SAME BASE. Law 4 needs the SAME EXPONENT. You cannot add exponents of 2^{3} × 3^{2} — the bases differ."),
]),
("The zero exponent — why n^{0} = 1", [
 ("p", "Walk down the powers of 2, dividing by 2 each step:"),
 ("eg", ["2^{3} = 8",
         "2^{2} = 4",
         "2^{1} = 2",
         "2^{0} = 1   ← keep dividing by 2, and 2 ÷ 2 = 1"]),
 ("key", "n^{0} = 1 for any n except 0. Not 0 — ONE."),
 ("p", "So 5^{0} + 6^{0} + 7^{0} = 1 + 1 + 1 = 3. Free marks, if you remember it."),
]),
("Negative exponents — just keep walking down", [
 ("p", "Do not stop at zero. Carry on dividing by 2:"),
 ("eg", ["2^{1} = 2",
         "2^{0} = 1",
         "2^{−1} = 1/2",
         "2^{−2} = 1/4",
         "2^{−3} = 1/8"]),
 ("key", "n^{−a} = 1/n^{a}.  A negative exponent means 'flip it', NOT 'make it negative'."),
 ("warn", "This is the single biggest trap in the chapter. 2^{−3} = 1/8. It is NOT −8, and it is NOT −6."),
]),
("Fractions with negative powers — the flip", [
 ("p", "Put those two ideas together and something neat falls out:"),
 ("eg", ["(2/3)^{−1} = 1 ÷ (2/3) = 3/2",
         "So a negative power simply TURNS THE FRACTION UPSIDE DOWN.",
         "(2/3)^{−3} = (3/2)^{3} = 27/8",
         "(5/8)^{−2} = (8/5)^{2} = 64/25"]),
 ("key", "(a/b)^{−n} = (b/a)^{n}.  Flip the fraction, and the power becomes positive."),
 ("p", "The MULTIPLICATIVE INVERSE (or reciprocal) of a number is what you multiply it by to get 1. Since n^{a} × n^{−a} = n^{0} = 1, the inverse of 5^{−2} is simply 5^{2} = 25."),
]),
("Solving equations with powers", [
 ("p", "Almost every 'find x' question in this chapter is the same one move."),
 ("steps", ["1. Rewrite BOTH sides so they have the SAME BASE.",
            "2. Once the bases match, the exponents must match too.",
            "3. Set the exponents equal and solve the little equation."]),
 ("eg", ["Find p if 12^{p} × 12^{−3} = (1/12)^{−5}",
         "Left side: same base, multiplying → add:  12^{p−3}",
         "Right side: flip the fraction → 12^{5}",
         "So 12^{p−3} = 12^{5}.  Bases match ⇒ p − 3 = 5 ⇒ p = 8"]),
 ("eg", ["Find x if 5^{−8} × (5^{3})^{x} = 25^{3} × 5^{x}",
         "Left: (5^{3})^{x} = 5^{3x}, so the left side is 5^{−8+3x}",
         "Right: 25 = 5^{2}, so 25^{3} = 5^{6}, and the right side is 5^{6+x}",
         "Bases match ⇒ −8 + 3x = 6 + x ⇒ 2x = 14 ⇒ x = 7"]),
 ("warn", "The commonest slip: forgetting to convert 25 into 5^{2}, or 81 into 3^{4}. If the bases are not the same, you may NOT compare exponents."),
]),
("Standard form — taming huge and tiny numbers", [
 ("p", "The mass of the Earth is about 6512000000000000000000000 kg. Nobody wants to count those zeros."),
 ("key", "Standard form (scientific notation): write the number as x × 10^{y}, where x is between 1 and 10, and y is a whole number."),
 ("steps", ["1. Put the decimal point just after the FIRST non-zero digit.",
            "2. Count how many places the point moved.",
            "3. Moved LEFT (big number) → positive power.  Moved RIGHT (small number) → negative power."]),
 ("eg", ["34500000 → put the point after the 3 → 3.45",
         "The point moved 7 places LEFT, so:  3.45 × 10^{7}"]),
 ("eg", ["0.00003412652 → put the point after the 3 → 3.412652",
         "The point moved 5 places RIGHT, so:  3.412652 × 10^{−5}"]),
 ("p", "Going backwards (into 'usual form') just undoes it: 3 × 10^{8} = 300000000."),
]),
("Adding and subtracting in standard form", [
 ("warn", "You CANNOT add 1.987 × 10^{11} and 6.54 × 10^{8} as they stand. The powers of 10 must match first — just like you cannot add halves and thirds until the denominators match."),
 ("steps", ["1. Rewrite the smaller number so it has the SAME power of 10 as the bigger one.",
            "2. Now add or subtract the front parts.",
            "3. Tidy the answer back into standard form (front part between 1 and 10)."]),
 ("eg", ["1.987 × 10^{11} − 6.54 × 10^{8}",
         "Make both 10^{8}:  1.987 × 10^{11} = 1987 × 10^{8}",
         "1987 × 10^{8} − 6.54 × 10^{8} = (1987 − 6.54) × 10^{8} = 1980.46 × 10^{8}",
         "Tidy up: 1980.46 × 10^{8} = 1.98046 × 10^{11}"]),
]),
("Multiplying, dividing and comparing", [
 ("p", "These are the EASY ones — the powers do not need to match. Just handle the front parts and the powers separately."),
 ("eg", ["(8.1 × 10^{16}) ÷ (3 × 10^{8})",
         "Front parts: 8.1 ÷ 3 = 2.7",
         "Powers: 10^{16} ÷ 10^{8} = 10^{16−8} = 10^{8}",
         "Answer: 2.7 × 10^{8}"]),
 ("eg", ["How many times heavier is the Earth (6.512 × 10^{24} kg) than the Moon (7.4 × 10^{22} kg)?",
         "Divide: (6.512 ÷ 7.4) × 10^{24−22} = 0.88 × 10^{2} = 88 times"]),
 ("warn", "When two quantities are given in different UNITS (km and m, say), convert them first. 3.844 × 10^{5} km is 3.844 × 10^{8} m — a thousand times bigger than it looks."),
]),
("⚠ Beyond your textbook — fractional exponents", [
 ("warn", "Your book's Summary (p.46) says the exponent is always a whole number. So this is extra — but one old paper uses it."),
 ("p", "A fractional power means a ROOT. The bottom of the fraction says which root; the top says which power."),
 ("eg", ["8^{1/3} = ∛8 = 2",
         "(8/27)^{2/3} = take the cube root, then square it = (2/3)^{2} = 4/9",
         "32^{−2/5} = flip, take the 5th root, then square = 1/(2^{2}) = 1/4"]),
]),
("Mistakes to avoid", [
 ("bullets", ["2^{−3} means 1/8. It does NOT mean −8.",
              "n^{0} = 1, not 0. Even (−2)^{0} = 1.",
              "A negative number to an EVEN power is positive: (−2)^{6} = +64. To an ODD power it stays negative: (−2)^{3} = −8.",
              "You may only add/subtract exponents when the BASES are the same. 2^{3} × 3^{2} does not simplify.",
              "Before comparing exponents, make the bases identical — rewrite 25 as 5^{2}, 81 as 3^{4}, 32 as 2^{5}.",
              "In standard form the front number must be between 1 and 10. 1980.46 × 10^{8} is not finished — it becomes 1.98046 × 10^{11}.",
              "To add or subtract in standard form, the powers of 10 must match FIRST."]),
]),
("Quick self-test", [
 ("check", [("Simplify 2^{5} × 2^{−3}", "2^{2} = 4  (same base, so add: 5 + (−3) = 2)"),
            ("What is (−3)^{0} + 3^{−2}?", "1 + 1/9 = 10/9"),
            ("Write (2/5)^{−3} without a negative power.", "(5/2)^{3} = 125/8  (flip the fraction)"),
            ("Find x:  3^{2x} = 81", "81 = 3^{4}, so 2x = 4 and x = 2"),
            ("Write 0.00047 in standard form.", "4.7 × 10^{−4}"),
            ("Simplify (10^{3})^{2} ÷ 10^{4}", "10^{6} ÷ 10^{4} = 10^{2} = 100"),
            ("(2.5 × 10^{6}) × (4 × 10^{3}) = ?", "10 × 10^{9} = 1 × 10^{10}  (tidy it — the front must be under 10)")]),
]),
]

# =========================== rendering ===========================
def build(fname, title, subtitle, sections, accent, tiermap, realmap):
    doc = Document()
    for s in doc.sections:
        s.top_margin = s.bottom_margin = Inches(0.7)
        s.left_margin = s.right_margin = Inches(0.8)
    doc.styles["Normal"].font.name = "Calibri"
    doc.styles["Normal"].font.size = Pt(10.5)

    t = doc.add_paragraph(); t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rich(t, title, size=20, bold=True, color=accent)
    s = doc.add_paragraph(); s.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rich(s, subtitle, size=12, color=NAVY)
    s2 = doc.add_paragraph(); s2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rich(s2, "Start at the top and read straight through. No prior knowledge assumed.", size=9.5, italic=True, color=GREY)
    s3 = doc.add_paragraph(); s3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rich(s3, "Every method here is one taught in NCERT Ganita Prakash, Grade 8.", size=9, italic=True, color=GREY)

    for i, (head, blocks) in enumerate(sections, 1):
        doc.add_paragraph()
        h = doc.add_paragraph()
        h.paragraph_format.space_after = Pt(1)
        rich(h, f"{i}.  {head}", size=13.5, bold=True, color=accent)
        if i in tiermap:
            tier, why = tiermap[i]
            bp = doc.add_paragraph()
            bp.paragraph_format.space_after = Pt(4)
            rich(bp, TIER_LBL[tier] + "  ", size=9, bold=True, color=TIER_COL[tier])
            rich(bp, why, size=9, italic=True, color=GREY)

        for kind, body in blocks:
            if kind == "p":
                p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(4)
                rich(p, body, size=10.5)
            elif kind == "key":
                p = doc.add_paragraph(); p.paragraph_format.left_indent = Inches(0.15)
                p.paragraph_format.space_after = Pt(5)
                rich(p, "KEY IDEA:  ", size=10, bold=True, color=GREEN)
                rich(p, body, size=10.5, bold=True, color=GREEN)
            elif kind == "warn":
                p = doc.add_paragraph(); p.paragraph_format.left_indent = Inches(0.15)
                p.paragraph_format.space_after = Pt(5)
                rich(p, "⚠  ", size=10, bold=True, color=AMBER)
                rich(p, body, size=10, italic=True, color=AMBER)
            elif kind == "eg":
                for j, line in enumerate(body):
                    p = doc.add_paragraph()
                    p.paragraph_format.left_indent = Inches(0.32)
                    p.paragraph_format.space_after = Pt(1)
                    rich(p, ("Example:  " if j == 0 else "") + line,
                         size=10, bold=(j == 0), color=BLUE)
                doc.add_paragraph().paragraph_format.space_after = Pt(3)
            elif kind == "steps":
                for line in body:
                    p = doc.add_paragraph()
                    p.paragraph_format.left_indent = Inches(0.32)
                    p.paragraph_format.space_after = Pt(1)
                    rich(p, line, size=10)
                doc.add_paragraph().paragraph_format.space_after = Pt(3)
            elif kind == "bullets":
                for line in body:
                    p = doc.add_paragraph(style="List Bullet")
                    p.paragraph_format.space_after = Pt(2)
                    rich(p, line, size=10)
            elif kind == "tbl":
                heads, rows = body
                tb = doc.add_table(rows=1, cols=len(heads))
                tb.style = "Light Grid Accent 1"
                tb.alignment = WD_TABLE_ALIGNMENT.CENTER
                for c, htxt in enumerate(heads):
                    cell = tb.rows[0].cells[c]
                    cell.paragraphs[0].text = ""
                    rich(cell.paragraphs[0], htxt, size=9.5, bold=True)
                for row in rows:
                    cells = tb.add_row().cells
                    for c, v in enumerate(row):
                        cells[c].paragraphs[0].text = ""
                        rich(cells[c].paragraphs[0], v, size=9.5)
                doc.add_paragraph().paragraph_format.space_after = Pt(3)
            elif kind == "check":
                for n_, (q, a) in enumerate(body, 1):
                    p = doc.add_paragraph()
                    p.paragraph_format.left_indent = Inches(0.15)
                    p.paragraph_format.space_after = Pt(1)
                    rich(p, f"Q{n_}. ", size=10, bold=True)
                    rich(p, q, size=10)
                    p2 = doc.add_paragraph()
                    p2.paragraph_format.left_indent = Inches(0.42)
                    p2.paragraph_format.space_after = Pt(4)
                    rich(p2, "\u2192 " + a, size=9.5, italic=True, color=GREEN)

        if i in realmap:
            rp = doc.add_paragraph()
            rp.paragraph_format.left_indent = Inches(0.15)
            rp.paragraph_format.space_before = Pt(3)
            rp.paragraph_format.space_after = Pt(6)
            rich(rp, "WHERE YOU ACTUALLY SEE THIS:  ", size=9.5, bold=True, color=TEAL)
            rich(rp, realmap[i], size=9.5, color=TEAL)

    doc.save(DIR + fname)
    print("saved:", fname, f"({len(sections)} sections)")


build("Maths - Chapter 1 Squares and Cubes - CONCEPTS EXPLAINED.docx",
      "Chapter 1 — A Square and A Cube",
      "Everything you need, explained from scratch",
      C1, NAVY, TIER1, REAL1)
build("Maths - Chapter 2 Power Play - CONCEPTS EXPLAINED.docx",
      "Chapter 2 — Power Play",
      "Exponents, powers and standard form, explained from scratch",
      C2, GREEN, TIER2, REAL2)
