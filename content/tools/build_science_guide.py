#!/usr/bin/env python3
"""A tough-test Science study guide, marked by how often each topic appears in
the real Sri Kumaran papers. Grounded in the NCERT 'Curiosity' Grade 8 chapters
(scoped exactly to the test) plus the fine-print boxes the school likes to ask.

Exam-frequency tiers (counted across 59 real Science question papers):
  Light: reflection & mirrors  30/59  ***
  Cell structure / organelles  29/59  ***  (of the Biology share)
  Microorganisms (OUT of scope)
  Light: lenses                15/59  **
  States of matter              2/59  *
"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

OUT = "/Users/shyamdk/Developer/personal/laya/data/question-papers/Science - Study Guide (Light, Cell, Matter).docx"

NAVY = RGBColor(0x1F, 0x3A, 0x5F)
GREY = RGBColor(0x60, 0x60, 0x60)
GREEN = RGBColor(0x1B, 0x5E, 0x20)
RED = RGBColor(0xC6, 0x28, 0x28)
AMBER = RGBColor(0xB2, 0x5E, 0x00)
TEAL = RGBColor(0x00, 0x69, 0x5C)
PURPLE = RGBColor(0x6A, 0x1B, 0x9A)

doc = Document()
for s in doc.sections:
    s.top_margin = s.bottom_margin = Inches(0.7)
    s.left_margin = s.right_margin = Inches(0.8)
st = doc.styles["Normal"]
st.font.name = "Calibri"
st.font.size = Pt(10.5)


def para(text="", size=10.5, bold=False, italic=False, color=None, after=4, align=None, indent=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(after)
    if align:
        p.alignment = align
    if indent:
        p.paragraph_format.left_indent = Inches(indent)
    if text:
        r = p.add_run(text)
        r.font.size, r.bold, r.italic = Pt(size), bold, italic
        if color:
            r.font.color.rgb = color
    return p


def rich(p, runs):
    for text, kw in runs:
        r = p.add_run(text)
        r.font.size = Pt(kw.get("size", 10.5))
        r.bold = kw.get("bold", False)
        r.italic = kw.get("italic", False)
        if "color" in kw:
            r.font.color.rgb = kw["color"]


def tier_badge(p, tier):
    color, label = {
        "***": (RED, "MUST KNOW"),
        "**": (AMBER, "IMPORTANT"),
        "*": (GREY, "SEEN SOMETIMES"),
    }[tier]
    r = p.add_run(f"{tier} {label}  ")
    r.bold = True
    r.font.size = Pt(9)
    r.font.color.rgb = color


def h1(text, accent=NAVY):
    doc.add_page_break()
    para(text, size=19, bold=True, color=accent, after=3)


def h2(text, tier=None, why=None, accent=NAVY):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(text)
    r.font.size, r.bold, r.font.color.rgb = Pt(13.5), True, accent
    if tier:
        pp = doc.add_paragraph()
        pp.paragraph_format.space_after = Pt(4)
        tier_badge(pp, tier)
        if why:
            rr = pp.add_run(why)
            rr.italic = True
            rr.font.size = Pt(9)
            rr.font.color.rgb = GREY


def box(label, text, color):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.12)
    p.paragraph_format.space_after = Pt(5)
    p.paragraph_format.space_before = Pt(2)
    r = p.add_run(f"{label}  ")
    r.bold = True
    r.font.size = Pt(9)
    r.font.color.rgb = color
    r2 = p.add_run(text)
    r2.font.size = Pt(10)
    r2.font.color.rgb = color


def bullets(items, color=None):
    for it in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(2)
        r = p.add_run(it)
        r.font.size = Pt(10)
        if color:
            r.font.color.rgb = color


def table(head, rows, widths=None):
    tb = doc.add_table(rows=1, cols=len(head))
    tb.style = "Light Grid Accent 1"
    tb.alignment = WD_TABLE_ALIGNMENT.CENTER
    for c, h in enumerate(head):
        cell = tb.rows[0].cells[c]
        cell.paragraphs[0].text = ""
        r = cell.paragraphs[0].add_run(h)
        r.bold = True
        r.font.size = Pt(9.5)
    for row in rows:
        cells = tb.add_row().cells
        for c, v in enumerate(row):
            cells[c].paragraphs[0].text = ""
            r = cells[c].paragraphs[0].add_run(v)
            r.font.size = Pt(9.5)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)


def qa(q, a):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(1)
    r = p.add_run("Q. ")
    r.bold = True
    r.font.size = Pt(10)
    r2 = p.add_run(q)
    r2.font.size = Pt(10)
    p2 = doc.add_paragraph()
    p2.paragraph_format.left_indent = Inches(0.28)
    p2.paragraph_format.space_after = Pt(5)
    r3 = p2.add_run("A.  " + a)
    r3.font.size = Pt(10)
    r3.font.color.rgb = GREEN


# ============================================================ COVER
para("Science — Study Guide", size=24, bold=True, color=NAVY, align=WD_ALIGN_PARAGRAPH.CENTER, after=2)
para("Grade 8 · for the coming test", size=13, color=NAVY, align=WD_ALIGN_PARAGRAPH.CENTER, after=2)
para("Physics: Light  ·  Biology: The Cell  ·  Chemistry: Particulate Nature of Matter",
     size=11, italic=True, color=GREY, align=WD_ALIGN_PARAGRAPH.CENTER, after=10)

para("What this guide covers, and how important each part is", size=12, bold=True, color=NAVY)
para("The importance stars are not a guess — they were counted across the 59 real Sri Kumaran "
     "Science papers from 2009 to 2026. Study the ✦✦✦ topics first.", size=9.5, italic=True, color=GREY, after=6)

table(["Topic", "In the papers", "Priority", "In your syllabus?"], [
    ["Physics — Light: reflection & mirrors", "30 of 59", "*** MUST KNOW", "Yes — whole chapter"],
    ["Biology — the cell & its organelles", "29 of 59", "*** MUST KNOW", "Yes"],
    ["Physics — Light: lenses", "15 of 59", "** IMPORTANT", "Yes — whole chapter"],
    ["Chemistry — states of matter", "2 of 59", "* SEEN SOMETIMES", "Yes — 7.1–7.2, up to p.106"],
    ["— microorganisms", "—", "excluded", "NO — not in the test"],
    ["— interparticle spacing (7.3+)", "—", "excluded", "NO — beyond p.106"],
])

box("HOW TO USE THIS", "Read the coloured boxes carefully — KEY FACT (green) is what you must be able to "
    "state, TRAP (amber) is where students lose marks, and FINE PRINT (teal) is the side-box "
    "detail the school likes to ask. Each chapter ends with a rapid-revision list and a self-test.", NAVY)

# ============================================================ PHYSICS
h1("1.  Physics — Light: Mirrors and Lenses", accent=PURPLE)
para("This is the highest-scoring topic in the whole Science paper — reflection and mirrors "
     "appeared in 30 of the 59 past papers. Master it completely.", size=10, italic=True, color=GREY, after=6)

h2("1.1  Reflection and the two laws", "***", "Reflection appears in 30 of the 59 papers.", PURPLE)
para("When light falls on a surface and bounces back, that is REFLECTION. To describe it we use three lines:")
bullets([
    "INCIDENT RAY — the ray of light falling on the mirror.",
    "REFLECTED RAY — the ray that comes back from the mirror.",
    "NORMAL — a line drawn at 90° (a right angle) to the mirror, at the exact point where the ray hits.",
])
box("KEY FACT", "The angle of incidence (i) and the angle of reflection (r) are ALWAYS measured "
    "from the NORMAL — never from the surface of the mirror.", GREEN)
para("The chapter gives two laws of reflection:", after=2)
para("First law:  the angle of incidence (i) is equal to the angle of reflection (r).   i = r",
     size=10.5, bold=True, indent=0.2, after=2)
para("Second law:  the incident ray, the normal, and the reflected ray all lie in the same plane.",
     size=10.5, bold=True, indent=0.2, after=4)
box("TRAP", "If a ray falls straight ALONG the normal, the angle of incidence is 0°, so the angle of "
    "reflection is also 0° — the ray reflects straight back on itself. (A common exam figure.)", AMBER)
box("TRAP", "The laws of reflection are valid for ALL mirrors — plane, concave AND convex — not just "
    "flat ones.", AMBER)

h2("1.2  Spherical mirrors: concave and convex", "***", "The single most-asked idea.", PURPLE)
para("A spherical mirror is a curved mirror — think of a small piece cut from a shiny ball.")
table(["", "Concave mirror", "Convex mirror"], [
    ["Curves…", "INWARDS (like a cave)", "OUTWARDS (bulges towards you)"],
    ["Parallel rays…", "converge (are brought together)", "diverge (are spread apart)"],
    ["Image of a NEAR object", "erect and enlarged", "erect and diminished"],
    ["Image of a FAR object", "inverted and diminished", "erect and diminished (always)"],
])
box("KEY FACT", "A CONVEX mirror ALWAYS gives an erect, diminished image, whatever the distance. "
    "A CONCAVE mirror changes: near object → erect & enlarged; far object → inverted & smaller.", GREEN)
box("KEY FACT", "A PLANE (flat) mirror always gives an erect image of the SAME size as the object. "
    "Only curved mirrors change the size.", GREEN)
para("Everyday uses (often asked):", size=10, bold=True, after=2)
table(["Concave mirror is used in…", "Convex mirror is used in…"], [
    ["torch & headlight reflectors", "vehicle side-view mirrors"],
    ["a dentist's mirror (enlarged view of teeth)", "road-safety mirrors at bends & junctions"],
    ["shaving/make-up mirrors", "surveillance mirrors in big stores"],
    ["solar cookers & the main mirror of a telescope", "—"],
])
box("TRAP", "Convex mirrors are chosen for vehicles and roads because they give a WIDER view of the "
    "traffic behind — not because the image is bigger (it is actually smaller).", AMBER)
box("TRAP", "'Lateral inversion' (left–right swap) happens in ALL three mirrors, including plane mirrors. "
    "Do not confuse it with the up–down 'inverted' image that only a concave mirror makes.", AMBER)

h2("1.3  Lenses: convex and concave", "**", "Lenses appeared in 15 of the 59 papers.", PURPLE)
para("A LENS is a piece of transparent material with at least one curved surface. Unlike a mirror, "
     "light passes THROUGH a lens instead of bouncing off it.")
table(["", "Convex lens", "Concave lens"], [
    ["Thicker at the…", "MIDDLE", "EDGES"],
    ["Parallel rays…", "converge", "diverge"],
    ["Image of a NEAR object", "erect and enlarged", "erect and diminished"],
    ["Image of a FAR object", "inverted and diminished", "erect and diminished (always)"],
    ["Common use", "magnifying glass", "—"],
])
box("KEY FACT", "Notice the pattern: a CONVEX LENS behaves like a CONCAVE MIRROR (both converge light; "
    "image depends on distance). A CONCAVE LENS behaves like a CONVEX MIRROR (both diverge; always "
    "erect & diminished).", GREEN)

h2("1.4  The one pattern that answers most image questions", "***", None, PURPLE)
box("KEY FACT", "The two that SPREAD light — CONVEX MIRROR and CONCAVE LENS — always give an ERECT, "
    "DIMINISHED image. The two that FOCUS light — CONCAVE MIRROR and CONVEX LENS — DEPEND on distance "
    "(far = inverted & small, near = erect & large). Learn this and you can answer almost any image "
    "question.", GREEN)

box("FINE PRINT", "A spherical mirror is NOT made by slicing a hollow glass sphere. It is made by "
    "grinding and polishing a flat glass piece into a curve. Coat the OUTER curved surface → a "
    "CONCAVE mirror; coat the INNER surface → a CONVEX mirror. (From the 'A step further' box — a "
    "classic tricky question.)", TEAL)
box("FINE PRINT", "About 800 years ago, in the time of the mathematician Bhāskara II, Indian astronomers "
    "used shallow bowls of water to observe stars and planets by reflection — showing they understood "
    "reflection in practice. Solar furnaces use concave mirrors to concentrate sunlight, hot enough to "
    "melt steel. (From 'Our scientific heritage'.)", TEAL)

para("Rapid revision — Light", size=11, bold=True, color=PURPLE, after=2)
bullets([
    "i = r, both measured from the NORMAL. Incident ray, normal, reflected ray are in one plane.",
    "Concave = curves in, converges. Convex = curves out, diverges.",
    "Convex mirror & concave lens: ALWAYS erect + diminished.",
    "Concave mirror & convex lens: far = inverted+small, near = erect+large.",
    "Plane mirror: erect, same size. Lateral inversion in all mirrors.",
    "Convex mirror → wider view (vehicles, roads, stores). Concave → focuses (torch, dentist, solar).",
    "Convex lens thicker in the middle (magnifying glass); concave lens thicker at the edges.",
    "Mirror made by coating flat glass: OUTER coat → concave, INNER coat → convex.",
])

para("Self-test — Light  (answers at the very end)", size=11, bold=True, color=PURPLE, after=3)
for i, q in enumerate([
    "A ray hits a mirror making 40° with the mirror surface. What is the angle of reflection? (careful!)",
    "Which mirror always gives an erect, diminished image?",
    "Why are convex mirrors used as side-view mirrors on vehicles?",
    "A lens is thicker at its edges than in the middle. Name it and say what it does to parallel light.",
    "Where is the reflective coating placed to make a concave mirror — the inner or outer surface?",
    "A concave mirror forms an inverted, smaller image. Is the object near the mirror or far from it?",
], 1):
    para(f"{i}. {q}", size=10, indent=0.12, after=2)

# ============================================================ BIOLOGY
h1("2.  Biology — The Cell and its Organelles", accent=GREEN)
para("Cell structure appeared in 29 of the 59 past papers. Note the SCOPE: only the cell and its "
     "parts are in your test — microorganisms (sections 2.3–2.4) are NOT.", size=10, italic=True, color=GREY, after=6)

h2("2.1  What is a cell?", "***", "Cell structure: 29 of 59 papers.", GREEN)
para("All living beings are made of CELLS — the smallest unit of life. Some organisms are a single "
     "cell (unicellular); larger living things, like us, are made of billions (multicellular).")
box("KEY FACT", "Every cell has THREE basic parts: the CELL MEMBRANE (outer boundary), the CYTOPLASM "
    "(the jelly-like filling), and the NUCLEUS (the round control centre). Plant cells have some extra "
    "parts on top of these.", GREEN)

h2("2.2  The organelles and what each one does", "***", "The core of the topic — learn every row.", GREEN)
table(["Part of the cell", "What it does", "Plant, animal, or both?"], [
    ["Cell membrane", "Encloses the cytoplasm & nucleus; separates one cell from the next; POROUS — lets useful things in and waste out.", "Both"],
    ["Cytoplasm", "Fills the space between membrane and nucleus; holds the other parts; MOST life processes happen here.", "Both"],
    ["Nucleus", "Regulates ALL the activities of the cell, and controls its growth.", "Both"],
    ["Cell wall", "An extra outer layer that gives rigidity and strength — so plant cells look firm and tightly packed.", "PLANT only"],
    ["Chloroplast (a plastid)", "Contains the green pigment chlorophyll and carries out photosynthesis. In non-green parts, plastids store food.", "PLANT only"],
    ["Vacuole", "Stores substances, removes waste, and keeps the cell's shape.", "LARGE in plant cells; small/absent in animal cells"],
])
box("TRAP", "The cell membrane is POROUS (it has tiny pores) — that is how materials move in and out. "
    "The CELL WALL, by contrast, is for strength. Do not swap their jobs.", AMBER)

h2("2.3  Plant cell vs animal cell", "**", "A very common comparison question.", GREEN)
box("KEY FACT", "A PLANT cell has three things an ANIMAL cell does NOT: a CELL WALL, CHLOROPLASTS, and "
    "a LARGE VACUOLE. Everything else (membrane, cytoplasm, nucleus) both cells share.", GREEN)
table(["Feature", "Plant cell", "Animal cell"], [
    ["Cell wall", "Present", "Absent"],
    ["Chloroplasts", "Present (green parts)", "Absent"],
    ["Vacuole", "One large vacuole", "Small, or none"],
    ["Shape", "Usually fixed / regular", "Usually rounder / irregular"],
])

h2("2.4  A cell's shape follows its job", "*", None, GREEN)
para("Cells are not all the same shape — the shape suits the work the cell does.")
bullets([
    "A MUSCLE cell is spindle-shaped (long and tapering) — suited to contracting.",
    "A NERVE cell is long with branches — suited to carrying signals over a distance.",
])

box("FINE PRINT", "In 1665, Robert Hooke looked at a thin slice of CORK through his microscope and saw "
    "tiny box-like rooms — he named them CELLS. He published his drawings in a book called MICROGRAPHIA. "
    "(This is a favourite exam question.)", TEAL)
box("TRAP", "This new textbook LABELS 'mitochondria' in the plant/animal cell diagram (Fig 2.5) but NEVER "
    "states its function. If your test asks what mitochondria does ('powerhouse of the cell'), that comes "
    "from your class notes, not this book. Check with your teacher.", AMBER)

para("Rapid revision — the Cell", size=11, bold=True, color=GREEN, after=2)
bullets([
    "Cell = smallest unit of life. Unicellular (one cell) vs multicellular (many).",
    "3 basic parts: cell membrane, cytoplasm, nucleus.",
    "Membrane = porous boundary; cytoplasm = filling where life processes happen; nucleus = control centre.",
    "Plant-only extras: cell wall (strength), chloroplast (photosynthesis), large vacuole.",
    "Animal cell: no wall, no chloroplast, only small vacuoles.",
    "Muscle cell = spindle-shaped; nerve cell = long & branched. Shape suits the job.",
    "Robert Hooke, 1665, cork → named 'cells', in the book Micrographia.",
])

para("Self-test — the Cell", size=11, bold=True, color=GREEN, after=3)
for i, q in enumerate([
    "Name the three basic parts found in every cell.",
    "Which part is porous and controls what enters and leaves the cell?",
    "List the three structures a plant cell has that an animal cell does not.",
    "Which organelle carries out photosynthesis, and what pigment does it contain?",
    "Who first observed and named cells, in which year, looking at what material?",
    "Why is a nerve cell long and branched?",
], 1):
    para(f"{i}. {q}", size=10, indent=0.12, after=2)

# ============================================================ CHEMISTRY
h1("3.  Chemistry — Particulate Nature of Matter", accent=NAVY)
para("SCOPE: only sections 7.1 and 7.2, up to book page 106. This topic is rarely asked directly "
     "(2 of 59 papers), but it is short and easy marks — and it explains why solids, liquids and gases "
     "behave as they do.", size=10, italic=True, color=GREY, after=6)

h2("3.1  What is matter made of?", "*", "Foundational; underpins the states of matter.", NAVY)
para("Break a stick of chalk, grind it to powder, look at it under a lens — every speck is still chalk. "
     "It never becomes a new substance.")
box("KEY FACT", "All matter is made of extremely small pieces called CONSTITUENT PARTICLES. Breaking, "
    "grinding or dissolving only makes the pieces smaller — it does not change the substance.", GREEN)
para("Dissolve sugar in water: it seems to disappear, but it has only broken into particles too small "
    "to see. The water still tastes sweet — the sugar is still there.")

h2("3.2  The three states of matter", "*", None, NAVY)
box("KEY FACT", "The constituent particles are held together by INTERPARTICLE ATTRACTIONS. The STRENGTH "
    "of that attraction decides whether a substance is a solid, a liquid, or a gas.", GREEN)
table(["", "Solid", "Liquid", "Gas"], [
    ["Attraction between particles", "strongest", "moderate", "weakest"],
    ["Particles", "tightly packed, fixed positions, only vibrate", "loosely packed, slide past each other", "far apart, move freely in all directions"],
    ["Shape", "definite (fixed)", "takes the container's shape", "no fixed shape — fills the container"],
    ["Volume", "definite", "definite", "no fixed volume"],
])
box("TRAP", "A LIQUID has a definite VOLUME but no definite SHAPE (it takes the shape of its container). "
    "Do not say a liquid has 'no fixed volume' — that is a gas.", AMBER)

h2("3.3  Changing state", "*", None, NAVY)
box("KEY FACT", "HEATING weakens the interparticle attraction, so the substance moves to a freer state:  "
    "SOLID  →(melting)→  LIQUID  →(boiling)→  GAS.", GREEN)
bullets([
    "MELTING — solid turning to liquid (at the melting point).",
    "BOILING — liquid turning to gas at the boiling point (e.g. water boils at 100 °C).",
    "EVAPORATION — a liquid slowly turning to gas even BELOW its boiling point.",
])
box("FINE PRINT", "Acharya Kanad, an ancient Indian philosopher, first spoke of the PARMANU (atom) — a "
    "tiny, indivisible, eternal particle of matter. He wrote this idea in his work, the VAISHESHIKA "
    "SUTRAS. (From 'Our scientific heritage' — commonly asked.)", TEAL)
box("TRAP", "Do NOT revise 'interparticle spacing' or the syringe/compression activity (section 7.3). "
    "It is beyond page 106 and NOT in your test.", AMBER)

para("Rapid revision — Matter", size=11, bold=True, color=NAVY, after=2)
bullets([
    "Matter = made of tiny constituent particles; breaking/grinding/dissolving doesn't change the substance.",
    "State is decided by the STRENGTH of interparticle attraction (solid > liquid > gas).",
    "Solid: definite shape & volume. Liquid: definite volume, no fixed shape. Gas: neither.",
    "Heating weakens attraction: solid →(melt)→ liquid →(boil)→ gas. Evaporation happens below boiling.",
    "Acharya Kanad → Parmanu (atom) → in the Vaisheshika Sutras.",
    "Scope stops at p.106 — interparticle spacing is OUT.",
])

para("Self-test — Matter", size=11, bold=True, color=NAVY, after=3)
for i, q in enumerate([
    "What decides whether a substance is a solid, a liquid or a gas?",
    "Which state has a definite volume but no definite shape?",
    "What happens to the interparticle attraction when a substance is heated?",
    "Name the process of a liquid becoming a gas below its boiling point.",
    "Which ancient Indian philosopher spoke of the Parmanu, and in which work?",
], 1):
    para(f"{i}. {q}", size=10, indent=0.12, after=2)

# ============================================================ ANSWERS
h1("Self-test answers", accent=GREEN)
para("Light", size=12, bold=True, color=PURPLE, after=2)
for a in [
    "50°. The 40° is measured from the SURFACE, so the angle from the normal is 90 − 40 = 50°. i = r = 50°.",
    "A convex mirror (and, for lenses, a concave lens).",
    "Because a convex mirror gives a WIDER field of view of the traffic behind.",
    "A concave lens; it diverges (spreads) parallel light.",
    "The OUTER curved surface is coated to make a concave mirror.",
    "The object is FAR from the mirror (a concave mirror only inverts when the object is far away).",
]:
    para("• " + a, size=10, color=GREEN, indent=0.1, after=2)
para("The Cell", size=12, bold=True, color=GREEN, after=2)
for a in [
    "Cell membrane, cytoplasm, and nucleus.",
    "The cell membrane.",
    "Cell wall, chloroplasts, and a large vacuole.",
    "The chloroplast; it contains chlorophyll.",
    "Robert Hooke, in 1665, looking at a slice of cork (published in Micrographia).",
    "So it can carry signals over a distance.",
]:
    para("• " + a, size=10, color=GREEN, indent=0.1, after=2)
para("Matter", size=12, bold=True, color=NAVY, after=2)
for a in [
    "The strength of the interparticle attraction.",
    "A liquid.",
    "It weakens.",
    "Evaporation.",
    "Acharya Kanad, in the Vaisheshika Sutras.",
]:
    para("• " + a, size=10, color=GREEN, indent=0.1, after=2)

doc.save(OUT)
print("saved:", OUT)
n_para = len(doc.paragraphs)
n_tbl = len(doc.tables)
print(f"  {n_para} paragraphs, {n_tbl} tables")
