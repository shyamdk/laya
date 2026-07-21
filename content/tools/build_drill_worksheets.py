#!/usr/bin/env python3
"""Printable Kumon-style speed & accuracy drill worksheets.

One Word document per strand (Addition, Subtraction, Multiplication Tables,
Multiplication, Division) — same "separate document per topic" convention as
the maths/science/social chapter docs. Each level gets WORKSHEETS_PER_LEVEL
fresh, non-duplicate worksheets (20 problems, 2 columns of 10 — the classic
Kumon layout), a self-check log to track her own times, and an answer key
section at the back so she can self-correct like the real method.

Mastery rule printed on the cover: redo the WHOLE worksheet (not just the
wrong ones) until she hits her own target time with at least 18/20 correct,
twice in a row — then move to the next level. Targets are calibrated from
HER first attempt, not a borrowed number, since Kumon's own published times
could not be verified (paywalled) and guessing them would be dishonest.
"""
import json
import os
import random

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from drill_gen import build_sheet

HERE = os.path.dirname(os.path.abspath(__file__))
LEVELS = json.load(open(os.path.join(HERE, "..", "data", "drill_levels.json")))
OUTDIR = os.path.join(HERE, "..", "..", "data", "drills")
os.makedirs(OUTDIR, exist_ok=True)

WORKSHEETS_PER_LEVEL = 3
PROBLEMS_PER_SHEET = 20  # 2 columns x 10 rows, matching the Kumon layout
SEED = 20260721  # fixed seed -> reproducible if regenerated

NAVY = RGBColor(0x1F, 0x3A, 0x5F)
GREY = RGBColor(0x6F, 0x71, 0x78)
INK = RGBColor(0x23, 0x26, 0x2C)
LINE = "D8D2C4"


def shade(cell, hexcolor):
    tcPr = cell._tc.get_or_add_tcPr()
    sh = OxmlElement("w:shd")
    sh.set(qn("w:val"), "clear")
    sh.set(qn("w:fill"), hexcolor)
    tcPr.append(sh)


def borders(tbl, hexcolor=LINE, sz=4):
    tblPr = tbl._tbl.tblPr
    b = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        e = OxmlElement(f"w:{edge}")
        e.set(qn("w:val"), "single")
        e.set(qn("w:sz"), str(sz))
        e.set(qn("w:color"), hexcolor)
        b.append(e)
    tblPr.append(b)


def new_doc():
    d = Document()
    for s in d.sections:
        s.top_margin = s.bottom_margin = Inches(0.6)
        s.left_margin = s.right_margin = Inches(0.75)
    d.styles["Normal"].font.name = "Calibri"
    d.styles["Normal"].font.size = Pt(11)
    return d


def para(doc, text="", size=11, bold=False, italic=False, color=None, after=4, align=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(after)
    if align:
        p.alignment = align
    if text:
        r = p.add_run(text)
        r.font.size, r.bold, r.italic = Pt(size), bold, italic
        if color:
            r.font.color.rgb = color
    return p


def cover(doc, strand_name, levels):
    para(doc, "Speed & Accuracy Drills", size=22, bold=True, color=NAVY, align=WD_ALIGN_PARAGRAPH.CENTER, after=2)
    para(doc, strand_name, size=15, color=NAVY, align=WD_ALIGN_PARAGRAPH.CENTER, after=14)

    para(doc, "How to use this booklet", size=12.5, bold=True, color=INK, after=4)
    for line in [
        "Time yourself on every worksheet — write the start and finish time in the boxes, "
        "or use a stopwatch, and work out the total.",
        "Check your own answers against the Answer Keys at the back as soon as you finish.",
        "Mastery = BOTH good speed AND good accuracy — not one or the other.",
        "Rule of thumb: your FIRST worksheet at a new level sets your baseline time. "
        "Aim to beat that baseline on your later attempts, with 18 or more out of 20 correct.",
        "If you don't hit your target time or miss more than 2, redo the WHOLE worksheet "
        "(the next one in this booklet) rather than just fixing the wrong ones — that's what "
        "builds real speed.",
        "Once you hit your target time with 18+/20 correct on two worksheets in a row, "
        "move to the next level.",
    ]:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(3)
        r = p.add_run(line)
        r.font.size = Pt(10)

    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    para(doc, "Levels in this booklet", size=12.5, bold=True, color=INK, after=4)
    t = doc.add_table(rows=1, cols=3)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    borders(t)
    hdr = t.rows[0].cells
    for c, h in enumerate(["Level", "Skill", "Worksheets"]):
        shade(hdr[c], "1F3A5F")
        r = hdr[c].paragraphs[0].add_run(h)
        r.bold, r.font.size, r.font.color.rgb = True, Pt(9.5), RGBColor(0xFF, 0xFF, 0xFF)
    for lvl in levels:
        cells = t.add_row().cells
        for c, v in enumerate([lvl["code"], lvl["title"], str(WORKSHEETS_PER_LEVEL)]):
            cells[c].paragraphs[0].text = ""
            r = cells[c].paragraphs[0].add_run(v)
            r.font.size = Pt(9.5)
    doc.add_page_break()


def log_box(doc):
    """A little self-tracking table: attempt / time / score, so progress is visible."""
    t = doc.add_table(rows=1, cols=4)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    borders(t, sz=3)
    hdr = t.rows[0].cells
    for c, h in enumerate(["Attempt", "Start", "Finish", "Score (out of 20)"]):
        shade(hdr[c], "F3EFE6")
        r = hdr[c].paragraphs[0].add_run(h)
        r.bold, r.font.size = True, Pt(8.5)
    for i in range(1, 4):
        cells = t.add_row().cells
        cells[0].paragraphs[0].add_run(str(i)).font.size = Pt(8.5)
        for c in (1, 2, 3):
            cells[c].paragraphs[0].add_run(" ").font.size = Pt(8.5)
    doc.add_paragraph().paragraph_format.space_after = Pt(3)


def worksheet_page(doc, strand_name, lvl, wnum, total, problems):
    para(doc, f"{strand_name} — {lvl['code']}", size=10, bold=True, color=GREY, after=1)
    para(doc, lvl["title"], size=15, bold=True, color=NAVY, after=6)
    para(doc, f"Worksheet {wnum} of {total}", size=10, italic=True, color=GREY, after=6)
    log_box(doc)

    half = len(problems) // 2
    left, right = problems[:half], problems[half:]
    t = doc.add_table(rows=half, cols=2)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    borders(t, sz=3)
    for r in range(half):
        for c, group in enumerate((left, right)):
            n = c * half + r + 1
            cell = t.rows[r].cells[c]
            cell.paragraphs[0].text = ""
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(3)
            p.paragraph_format.space_after = Pt(3)
            run = p.add_run(f"{n:>2}.  {group[r]['stem']}   ")
            run.font.size = Pt(12)
            blank = p.add_run("_______")
            blank.font.size = Pt(12)
    doc.add_page_break()


def answer_key(doc, strand_name, sheets):
    """sheets: list of (level, wnum, problems)"""
    para(doc, "Answer Keys", size=18, bold=True, color=NAVY, after=8)
    para(doc, "Self-check as soon as you finish a worksheet.", size=10, italic=True, color=GREY, after=10)
    for lvl_code, lvl_title, wnum, problems in sheets:
        para(doc, f"{lvl_code} — {lvl_title} · Worksheet {wnum}", size=11, bold=True, color=NAVY, after=2)
        half = len(problems) // 2
        left, right = problems[:half], problems[half:]
        t = doc.add_table(rows=half, cols=2)
        t.alignment = WD_TABLE_ALIGNMENT.CENTER
        borders(t, sz=2)
        for r in range(half):
            for c, group in enumerate((left, right)):
                n = c * half + r + 1
                cell = t.rows[r].cells[c]
                cell.paragraphs[0].text = ""
                run = cell.paragraphs[0].add_run(f"{n:>2}.  {group[r]['answer']}")
                run.font.size = Pt(9.5)
        doc.add_paragraph().paragraph_format.space_after = Pt(6)


def build_strand(strand):
    rng = random.Random(SEED + hash(strand["code"]) % 10000)
    doc = new_doc()
    cover(doc, strand["name"], strand["levels"])

    key_sheets = []
    for lvl in strand["levels"]:
        for w in range(1, WORKSHEETS_PER_LEVEL + 1):
            problems = build_sheet(lvl["gen"], PROBLEMS_PER_SHEET, rng)
            worksheet_page(doc, strand["name"], lvl, w, WORKSHEETS_PER_LEVEL, problems)
            key_sheets.append((lvl["code"], lvl["title"], w, problems))

    answer_key(doc, strand["name"], key_sheets)

    fname = f"Drills - {strand['name']}.docx"
    path = os.path.join(OUTDIR, fname)
    doc.save(path)
    n_sheets = len(strand["levels"]) * WORKSHEETS_PER_LEVEL
    print(f"  {strand['name']:<28} {len(strand['levels'])} levels x {WORKSHEETS_PER_LEVEL} sheets "
          f"= {n_sheets} worksheets  ->  {path}")


if __name__ == "__main__":
    print("Building drill worksheet booklets...")
    for strand in LEVELS["strands"]:
        build_strand(strand)
    print("\ndone.")
