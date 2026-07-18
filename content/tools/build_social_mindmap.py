#!/usr/bin/env python3
"""Printable mind-map version of the three Social chapters.

Word can't draw a radial map, so each chapter is a shaded central node
followed by its branches as colour + star coded boxes (left rail = tier).
Same importance scheme as the study guide: *** must know, ** important,
* good to know. History is flagged new-syllabus (tiered by emphasis).
"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUT = "/Users/shyamdk/Developer/personal/laya/data/question-papers/Social Science - Mind Maps (Resources, History, Civics).docx"

GEO = RGBColor(0x1F, 0x7A, 0x3F)
HIST = RGBColor(0x9A, 0x5A, 0x06)
CIV = RGBColor(0x27, 0x4F, 0x7D)
STAR = RGBColor(0xC0, 0x56, 0x2A)
GREY = RGBColor(0x6F, 0x71, 0x78)
DIM = RGBColor(0xB6, 0xB0, 0xA4)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
INK = RGBColor(0x23, 0x26, 0x2C)

# soft tint hexes (cell shading) per accent
TINT = {"geo": "E7F1EA", "hist": "F3E9DA", "civ": "E4ECF4"}
ACCENT = {"geo": GEO, "hist": HIST, "civ": CIV}

doc = Document()
for s in doc.sections:
    s.top_margin = s.bottom_margin = Inches(0.6)
    s.left_margin = s.right_margin = Inches(0.7)
doc.styles["Normal"].font.name = "Calibri"
doc.styles["Normal"].font.size = Pt(10.5)


def shade(cell, hexcolor):
    tcPr = cell._tc.get_or_add_tcPr()
    sh = OxmlElement("w:shd")
    sh.set(qn("w:val"), "clear")
    sh.set(qn("w:fill"), hexcolor)
    tcPr.append(sh)


def no_borders(tbl):
    tblPr = tbl._tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        e = OxmlElement(f"w:{edge}")
        e.set(qn("w:val"), "none")
        borders.append(e)
    tblPr.append(borders)


def stars(tier):
    """returns (full, empty) counts for a tier 1..3"""
    return tier, 3 - tier


def run(p, text, size=10.5, bold=False, italic=False, color=None):
    r = p.add_run(text)
    r.font.size, r.bold, r.italic = Pt(size), bold, italic
    if color:
        r.font.color.rgb = color
    return r


def star_runs(p, tier, on_color=STAR, size=10):
    full, empty = stars(tier)
    run(p, "★" * full, size=size, color=on_color)
    if empty:
        run(p, "★" * empty, size=size, color=DIM)


def para(text="", size=10.5, bold=False, italic=False, color=None, after=4, align=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(after)
    if align:
        p.alignment = align
    if text:
        run(p, text, size, bold, italic, color)
    return p


def hub(key, area, title, freq, new=False):
    """central node: full-width shaded box."""
    t = doc.add_table(rows=1, cols=1)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    no_borders(t)
    cell = t.rows[0].cells[0]
    shade(cell, {"geo": "1F7A3F", "hist": "9A5A06", "civ": "274F7D"}[key])
    cell.width = Inches(7.0)
    p1 = cell.paragraphs[0]
    p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p1.paragraph_format.space_after = Pt(1)
    run(p1, area.upper(), size=8, bold=True, color=WHITE)
    p2 = cell.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.paragraph_format.space_after = Pt(2)
    run(p2, title, size=15, bold=True, color=WHITE)
    p3 = cell.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p3.paragraph_format.space_after = Pt(1)
    if new:
        run(p3, "NEW syllabus", size=9, bold=True, color=WHITE)
    else:
        star_runs(p3, freq[0], on_color=WHITE, size=9)
        run(p3, f"  ·  {freq[1]}", size=9, bold=True, color=WHITE)
    doc.add_paragraph().paragraph_format.space_after = Pt(3)


def branch(key, tier, title, leaves):
    """one branch: 2-col row — tier rail + title | leaf bullets."""
    t = doc.add_table(rows=1, cols=2)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    no_borders(t)
    t.columns[0].width = Inches(2.35)
    t.columns[1].width = Inches(4.65)
    left, right = t.rows[0].cells
    left.width = Inches(2.35)
    right.width = Inches(4.65)
    shade(left, TINT[key])
    lp = left.paragraphs[0]
    lp.paragraph_format.space_after = Pt(1)
    run(lp, title, size=10.5, bold=True, color=ACCENT[key])
    sp = left.add_paragraph()
    sp.paragraph_format.space_after = Pt(1)
    star_runs(sp, tier)
    first = right.paragraphs[0]
    for i, lf in enumerate(leaves):
        p = first if i == 0 else right.add_paragraph()
        p.paragraph_format.space_after = Pt(1)
        run(p, "›  ", size=9.5, bold=True, color=ACCENT[key])
        # bold marker: text may contain **term**
        parts = lf.split("**")
        for j, seg in enumerate(parts):
            if seg:
                run(p, seg, size=9.5, bold=(j % 2 == 1), color=INK if j % 2 == 1 else GREY)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def note(key, tag, text):
    t = doc.add_table(rows=1, cols=1)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    no_borders(t)
    cell = t.rows[0].cells[0]
    shade(cell, TINT[key])
    cell.width = Inches(7.0)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(1)
    run(p, f"{tag}   ", size=8.5, bold=True, color=ACCENT[key])
    parts = text.split("**")
    for j, seg in enumerate(parts):
        if seg:
            run(p, seg, size=9.5, bold=(j % 2 == 1), color=INK)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)


# ============================================================ COVER
para("Social Science — Chapter Mind Maps", size=22, bold=True, color=INK,
     align=WD_ALIGN_PARAGRAPH.CENTER, after=2)
para("Grade 8 · the big picture, one chapter at a time", size=11, color=GREY,
     align=WD_ALIGN_PARAGRAPH.CENTER, after=8)

lg = doc.add_paragraph()
lg.alignment = WD_ALIGN_PARAGRAPH.CENTER
lg.paragraph_format.space_after = Pt(2)
run(lg, "How to read the stars:   ", size=9.5, bold=True, color=INK)
star_runs(lg, 3); run(lg, " must know      ", size=9.5, color=GREY)
star_runs(lg, 2); run(lg, " important      ", size=9.5, color=GREY)
star_runs(lg, 1); run(lg, " good to know", size=9.5, color=GREY)
para("The bigger a branch leans on the exam, the more stars — counted from the 51 past papers "
     "(except History; see its note).", size=9, italic=True, color=GREY,
     align=WD_ALIGN_PARAGRAPH.CENTER, after=8)

# ============================================================ CH1 — GEOGRAPHY
para("Chapter 1", size=10, bold=True, color=GEO, after=1)
hub("geo", "Geography", "Natural Resources & their Uses", (3, "32 / 51 papers"))
branch("geo", 1, "What is a resource?",
       ["Occurs in **Nature** and is **useful to humans**",
        "Trace anything back far enough — it leads to Nature"])
branch("geo", 3, "Grouped by USE",
       ["**Essential for life** — air, water, soil / food",
        "**For materials** — wood, marble, metals",
        "**For energy** — coal, water, petroleum, sun, wind"])
branch("geo", 3, "Renewable vs non-renewable",
       ["**Renewable** — sun, wind, water, forests (regrow if not overused)",
        "**Non-renewable** — coal, petroleum, minerals (run out)"])
branch("geo", 2, "Spread across India",
       ["Uneven — iron ore, coal, oil (Mumbai High)",
        "**Paradox of plenty** — resource-rich, grows slower",
        "Example: Punjab depleting its groundwater"])
branch("geo", 3, "Wise use — stewardship",
       ["Using resources so the **future** still has them",
        "**Sikkim** — fully organic state",
        "**Raichur** — solar energy; rainwater harvesting",
        "Fine print: Nature seen as **sacred**"])

doc.add_page_break()

# ============================================================ CH2 — HISTORY
para("Chapter 2", size=10, bold=True, color=HIST, after=1)
hub("hist", "History", "Reshaping India's Political Map", None, new=True)
branch("hist", 2, "The timeline",
       ["**1206** Delhi Sultanate begins",
        "**1526** Panipat I — Babur founds the Mughals",
        "**1556** Panipat II — Akbar",
        "**1565** Talikota · **1671** Saraighat · **1699** Khalsa"])
branch("hist", 2, "Delhi Sultanate",
       ["Dynasties ending with the **Lodis**",
        "Khilji — pushed south, repelled the Mongols",
        "**Muhammad bin Tughlaq** → capital to Daulatabad (failed)",
        "**Jizya** — tax on non-Muslim subjects"])
branch("hist", 2, "Vijayanagara",
       ["Founded by **Harihara & Bukka**; capital **Hampi**",
        "Peak under **Krishnadeva Raya**",
        "Fell at the Battle of **Talikota** (1565)"])
branch("hist", 2, "The Mughals",
       ["**Babur** — Panipat I (1526), wrote the Baburnama",
        "**Akbar** — Panipat II, Abul Fazl, **abolished jizya**",
        "Sher Shah Suri — the Sur Empire interlude"])
branch("hist", 2, "Resistance",
       ["**Rajputs** — Maharana Pratap, Haldighati (1576)",
        "**Ahoms** — Battle of Saraighat (1671)",
        "**Sikhs** — Khalsa (1699); Guru Tegh Bahadur"])
branch("hist", 1, "Fine print",
       ["**Musunuri Nayakas** — confederacy vs the Sultanate",
        "**Hundi** — a trade payment note"])
note("hist", "READ THIS",
     "**This chapter is new to the 2026-27 syllabus**, so the old papers can't rank it — don't trust "
     "a 'low frequency' here. Every branch matters; learn the timeline cold, since the rest hangs off "
     "those six dates.")

doc.add_page_break()

# ============================================================ CH3 — CIVICS
para("Chapter 3", size=10, bold=True, color=CIV, after=1)
hub("civ", "Civics", "Universal Franchise & Elections", (2, "17 / 51 papers"))
branch("civ", 2, "Universal adult franchise",
       ["Every adult **18+** gets **one vote**",
        "Regardless of caste, religion, gender, wealth, education",
        "Guaranteed by **Article 326**"])
branch("civ", 2, "Election Commission",
       ["ECI runs national, state & President polls",
        "Headed by the **Chief Election Commissioner**",
        "Officers: CEO → DEO → Returning Officer → ERO",
        "**T.N. Seshan** (1990) — brought in voter ID cards"])
branch("civ", 1, "Voting tools",
       ["**EVM** + **VVPAT** (paper trail to verify your vote)",
        "**NOTA** — reject all candidates",
        "**Model Code of Conduct** — fair-play rules"])
branch("civ", 2, "Lok Sabha & Rajya Sabha",
       ["**Lok Sabha** — 543 seats, direct, First-Past-the-Post",
        "**Rajya Sabha** — indirect, 6-yr term, 12 nominated",
        "**President / VP** — chosen by an electoral college"])
branch("civ", 1, "Fine print",
       ["**Uttaramerur** inscriptions — choosing reps by drawing lots",
        "~3.1 million elected representatives in local bodies"])
note("civ", "WATCH",
     "**Don't swap the two houses:** Lok Sabha is **directly** elected; Rajya Sabha is **indirect**. "
     "And First-Past-the-Post means the winner needs the most votes, not a majority.")

doc.save(OUT)
print("saved:", OUT)
print(f"  {len(doc.paragraphs)} paragraphs, {len(doc.tables)} tables")
