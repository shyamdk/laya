#!/usr/bin/env python3
"""Build the Class 8 Maths question bank (Ch1 Squares & Cubes, Ch2 Power Play)
from Sri Kumaran past papers 2021-22 .. 2025-26."""
import re
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

OUT = "/Users/shyamdk/Developer/personal/laya/data/question-papers/Maths - Squares Cubes & Power Play (2021-2026).docx"

# ---------------------------------------------------------------- content
# (question text, source, marks)   ^{..} = superscript, _{..} = subscript
CH1 = [
 ("Which of the following is a perfect square?\n(a) 6250 (b) 10404 (c) 38943 (d) 12547", "2025-26 Half Yearly", 1),
 ("If ∛636056 = 86, then (∛636.056 × ∛636056 × ∛0.000636056) is\n(a) 63.6056 (b) 636.056 (c) 0.636056 (d) 6.36056", "2025-26 Half Yearly", 1),
 ("If the cube root of x is 8, then the cube root of 8x is\n(a) 4 (b) 8 (c) 16 (d) 2", "2025-26 Annual", 1),
 ("Which of the following is a perfect square number?\n(a) 2^{6} × 3^{3} (b) 2^{5} × 7^{2} (c) 3 × 2 × 7^{2} (d) 11 × 3^{3} × 33", "2024-25 Half Yearly", 1),
 ("The least number by which 486 should be divided to make it a perfect square is\n(a) 6 (b) 3 (c) 5 (d) 7", "2024-25 Annual", 1),
 ("Which of the following statements is true?\n(a) There are 5 perfect cubes between 1 and 100.\n(b) If x^{2} ends with 5, then x^{3} also ends with 5.\n(c) Cube of a one-digit number cannot be a two-digit number.\n(d) Cube root of 8 is +2 and −2.", "2024-25 Annual", 1),
 ("If √11025 = x, then √1.1025 = ?\n(a) x/10 (b) x (c) x/100 (d) x/1000", "2023-24 Final", 1),
 ("A square number has 5 digits. Its square root is\n(a) always a 2-digit number (b) always a 2-digit number\n(c) always a 3-digit number (d) could be a 2- or a 3-digit number", "2022-23 Half Yearly", 1),
 ("Cube of 0.3 is\n(a) 0.27 (b) 2.7 (c) 0.0027 (d) 0.027", "2022-23 Annual", 1),
 ("25% of which of the following perfect cube numbers is equal to the cube root of that number?\n(a) 125 (b) 64 (c) 8 (d) 343", "2021-22 Annual", 1),
 ("If one member of a Pythagorean triplet is 12, which of these could be the remaining members?\n(a) 35 and 6 (b) 35 and 37 (c) 37 and 36 (d) 36 and 6", "2021-22 PT-2", 1),
 ("The prime factorisation of a number is 2 × 3 × 5 × 2^{5} × 5^{2}. To make the number a perfect square we can\n (i) multiply the number by 2 × 3 × 5 (ii) multiply the number by 5 × 3\n (iii) divide the number by 5 × 3 (iv) divide the number by 2 × 3 × 5\n(a) (i) and (iv) (b) (i) and (iii) (c) (ii) and (iii) (d) (ii) and (iv)", "2021-22 PT-2", 1),
 ("Which digit will be in the unit's place of the square of 3457?", "2024-25 PT-2", 1),
 ("The value of √(128 − √49) is ________.", "2022-23 Annual", 1),
 ("Using estimation, the cube root of 74088 is ________.", "2022-23 Annual", 1),
 ("Evaluate: (141)^{2} − (140)^{2}", "2025-26 Half Yearly", 1),
 ("Fill in the boxes with the correct numbers:\n  ___ + 23 + ___ + 27 + ___ = (___)^{3}", "2025-26 Half Yearly", 1),
 ("A cubical vessel can hold 9261 cubic cm of water when filled to the brim. What is its height?", "2025-26 Annual", 1),
 ("Estimate the cube root of 103823.", "2025-26 Annual", 1),
 ("Estimate the length of each side of a cube if its volume is 12167 cm^{3}.", "2024-25 Half Yearly", 1),
 ("If √7396 = 86, the value of √73960000 + √0.007396 is ________.", "2024-25 Half Yearly", 1),
 ("Estimate the cube root of 13824.", "2023-24 Final", 1),
 ("Which of the following is definitely NOT a square number? Justify with a suitable reason.\n (a) 1253687  (b) 2318025", "2022-23 Half Yearly", 1),
 ("The units digit and the tens digit of ∛658503 are ________ and ________ respectively.", "2023-24 PT-2", 1),
 ("Find 23^{2} − 22^{2} using a property of square numbers.", "2023-24 PT-2", 1),
 ("Find the Pythagorean triplet whose largest member is 26.", "2023-24 PT-2", 1),
 ("If the hypotenuse of a right-angled triangle is 17, find the lengths of the other two sides.", "2024-25 PT-2", 2),
 ("If one member of a Pythagorean triplet is 65, find the other two members.", "2025-26 Half Yearly", 2),
 ("Find the cube root of (i) 79507 (using the shortcut method)  (ii) 27/3^{9}", "2025-26 Half Yearly", 2),
 ("Find the value of x, if 1323/x is a perfect cube.", "2025-26 Half Yearly", 2),
 ("If x = ∛(27/64) and y = ∛(−1/64), what is x + y?", "2024-25 Half Yearly", 2),
 ("What is the smallest number by which 8820 must be multiplied to make it a perfect square? Also find the square root of the number so obtained.", "2024-25 PT-2", 3),
 ("Find the smallest number which must be subtracted from 5676 to make it a perfect square. Also find the square root of the perfect square so obtained.", "2023-24 PT-2", 2),
 ("Simplify: (∛343 + ∛0.343) × ∛10^{6}", "2023-24 PT-2", 2),
 ("Find the smallest number by which 2352 must be multiplied to obtain a perfect square. Find the square root of the new number so formed.", "2022-23 Half Yearly", 2),
 ("Find the value of √625 + √0.000625 + √0.0625", "2022-23 Half Yearly", 2),
 ("Find the cube root of (18)^{3} × (1/12)^{3}.", "2021-22 Annual", 1),
 ("Find the cube root of 19683 by estimation.", "2021-22 Annual", 2),
 ("If √3844 = 62 and √32.49 = 5.7, find the value of √0.3844 − √0.003249", "2021-22 PT-2", 2),
 ("Find the cube root of 17576 by prime factorisation.", "2024-25 Annual", 2),
 ("Find √20 up to two decimal places.", "2024-25 Annual", 3),
 ("Using long division, find √691.69", "2025-26 Half Yearly", 3),
 ("Find the square root of 125.44 using the long division method.", "2025-26 Annual", 3),
 ("Find the square root of 2079.36.", "2024-25 Half Yearly", 3),
 ("Find the square root of 108.16 by the long division method.", "2023-24 PT-2", 4),
 ("Find √54.76 using the long division method.", "2023-24 Final", 2),
 ("Find the square root of 8136.04 by long division.", "2022-23 Half Yearly", 3),
 ("Find the smallest square number which is divisible by 10, 20 and 28.", "2025-26 Half Yearly", 3),
 ("Is 2560 a perfect cube? If not, find the smallest natural number by which 2560 should be multiplied so that the product is a perfect cube.", "2025-26 Half Yearly", 3),
 ("Find the smallest perfect cube that is divisible by 4, 9 and 10.", "2024-25 Half Yearly", 3),
 ("Find the smallest perfect square number which is exactly divisible by 80, 12, 24 and 36. Also find its square root.", "2022-23 Half Yearly", 3),
 ("Find the smallest whole number by which 23625 must be divided so that the quotient is a perfect cube. Also find the cube root of the cube number so obtained.", "2022-23 Annual", 3),
 ("Find (a) the least number which must be subtracted from 7400 so as to get a perfect square, and (b) the least number which must be added to 7400 so as to get a perfect square.", "2021-22 PT-2", 3),
 ("(a) Between 23^{2} and 24^{2} there are ________ natural numbers.\n(b) The square root of the number 161728401 will have ________ digits.\n(c) The sum of the first 17 odd natural numbers is ________.", "2024-25 Half Yearly", 3),
 ("(a) Write the Pythagorean triplet whose one of the numbers is 24.\n(b) How many natural numbers lie between 18^{2} and 19^{2}?\n(c) Express 36 as the sum of consecutive odd numbers.", "2022-23 Half Yearly", 4),
 ("Sakshi makes a cuboid of sides 15 cm × 30 cm × 15 cm. How many such cuboids will she need to form a cube?", "2023-24 Final", 2),
 ("Three friends had some money in the ratio 2 : 3 : 4. If the sum of the cubes of the amounts is ₹792, find how much money each of them had.", "2025-26 Half Yearly", 3),
 ("Three numbers are in the ratio 2 : 3 : 4. If the sum of their cubes is 131769, find the three numbers.", "2023-24 PT-2", 3),
 ("As part of the Independence Day programme, a group of people were arranged so that the number of rows was equal to the number of columns, forming a square formation. If there were 9450 people, how many more persons are required to make this arrangement possible? Also find the number of persons in each row.", "2025-26 Half Yearly", 4),
 ("For the Republic Day parade, dancers are to be arranged in a formation such that the number of rows equals the number of columns. There are 2016 dancers in the group.\n(a) What is the minimum number of dancers to be added to make this arrangement?\n(b) Find the number of dancers in each row.", "2024-25 Half Yearly", 3),
 ("Daisy is decorating flowers in rows such that the number of flowers in one row is equal to the number of rows. She has 5646 flowers. How many more flowers does she need to make such an arrangement? How many flowers does she place in a row?", "2022-23 Half Yearly", 4),
 ("In a meeting hall there are 13695 chairs. The organiser arranged the chairs so that the number of chairs in each row is equal to the number of rows. Find the least number of chairs left out of this arrangement.", "2022-23 Annual", 3),
 ("The area of a square ground is 7056 m^{2}. Find the perimeter of the square ground. Also find the cost of fencing it at the rate of ₹12 per metre.", "2024-25 Half Yearly", 4),
]

CH2 = [
 ("(−3p^{−3})^{2} = ________\n(a) −3p^{−6} (b) 9p^{−6} (c) −9p^{6} (d) 6p^{6}", "2025-26 Annual", 1),
 ("The value of (0.03)^{−2} × (0.000027)^{3} is\n(a) (0.03)^{5} (b) (0.03)^{7} (c) 0.03 (d) (0.3)^{7}", "2025-26 Half Yearly", 1),
 ("The value of [ (2/3)^{2} × (−1/3)^{−2} ]^{2} is ________\n(a) 4^{2} (b) (1/4)^{2} (c) (4/81)^{2} (d) (81/4)^{2}", "2024-25 Half Yearly", 1),
 ("Which of these is equal to (10)^{−6}?\n(a) 10^{−3} ÷ 10^{9} (b) 10^{3} ÷ 10^{−9} (c) 10^{3} ÷ 10^{9} (d) 10^{−3} ÷ 10^{−9}", "2023-24 Half Yearly", 1),
 ("What number comes in the blank to make the number sentence true?\n (7/−5)^{3} × ________ = 1\n(a) (−5/7)^{−3} (b) (−7/5)^{−3} (c) (5/7)^{3} (d) (7/5)^{3}", "2022-23 Half Yearly", 1),
 ("The value of 1^{−7} + 1^{7} + 1^{0} is\n(a) 0 (b) 1 (c) 2 (d) 3", "2022-23 PT-1", 1),
 ("The value of x^{−1} ÷ y^{−1} is\n(a) y/x (b) x/y (c) 1/xy (d) xy/1", "2022-23 PT-1", 1),
 ("(12^{2} − 5^{3}) × (−1)^{20}/19 equals ________\n(a) 0 (b) 1 (c) −1 (d) 2", "2021-22 Half Yearly", 1),
 ("If x = (5/8)^{−2} × (4/5)^{−2}, then the value of x^{−1} is ________\n(a) 1/4 (b) 64 (c) 4 (d) 1/64", "2021-22 Half Yearly", 1),
 ("Which of the following statements is true?\n(a) 4.26 × 10^{6} = 4260 × 10^{4} (b) 4.26 × 10^{5} = 42600000\n(c) 4.26 × 10^{7} = 42600000 (d) 4.26 × 10^{6} = 426 × 10^{5}", "2021-22 PT-1", 1),
 ("The multiplicative inverse of 5^{−2} is ________\n(a) 1 (b) 5 (c) 10 (d) 25", "2021-22 PT-1", 1),
 ("The value of (−2)^{0} − (−2)^{1} + (−2)^{2} + (−2)^{3} is ________.", "2023-24 Half Yearly", 1),
 ("Write 0.00003412652 in standard form.", "2025-26 Half Yearly", 1),
 ("Express 34500000 in standard form.", "2021-22 PT-1", 1),
 ("Fill in the blank: ________ ÷ (1/5)^{−3} = (1/5)^{2}", "2024-25 Half Yearly", 1),
 ("Find the value of (1^{−1} − 3^{−1})^{−1}.", "2025-26 Annual", 1),
 ("Two steel rods weigh 15 × 10^{3} kg and 2.6 × 10^{2} kg respectively. Find the total weight of the rods.", "2024-25 Annual", 1),
 ("If (5x^{a})(bx^{2}) = 35x^{8}, find the values of a and b (where a and b are natural numbers).", "2024-25 Annual", 1),
 ("Find the value of 5^{−1}(5^{0} + 6^{0} + 7^{0}) + 5^{−1}.", "2022-23 Half Yearly", 1),
 ("Simplify: (24/3)^{4} × (3/8)^{4}", "2021-22 PT-1", 2),
 ("Simplify: 2^{−5} − 8^{−2} + 2^{−3}", "2023-24 Half Yearly", 2),
 ("Evaluate: [ (−2/3)^{3} × (−2/3) ] ÷ (4/9)^{2}", "2022-23 Half Yearly", 2),
 ("Evaluate: [ (5^{2})^{−1} × 125 ] ÷ (1/5)^{−2}", "2024-25 Annual", 2),
 ("Simplify: (p^{7} q^{2} r^{−4}) / (r^{−8} p^{3} q^{−2})", "2025-26 Half Yearly", 2),
 ("Find the value of x:  5^{−8} × (5^{3})^{x} = 25^{3} × 5^{x}", "2025-26 Half Yearly", 2),
 ("Find the value of p if  12^{p} × 12^{−3} = (1/12)^{−5}", "2023-24 Half Yearly", 2),
 ("Find the value of x if  (4/5)^{3} ÷ (5/4)^{3} = (4/5)^{3x}", "2022-23 PT-1", 2),
 ("Find the value of x:  [ (2/5)^{−2} ]^{2x} = 625/16", "2022-23 Half Yearly", 2),
 ("If a/b = (2/3)^{−3} × (5/6)^{0}, find the value of (a/b)^{−2}.", "2021-22 Half Yearly", 2),
 ("Find x, if  (5/9)^{2x} × (5/9)^{−18} = (81/25)^{−3}", "2025-26 Annual", 2),
 ("Find the value of x using the laws of exponents:\n [ (9/16)^{4} × (4/3)^{10} ] ÷ (64/27) = (4/3)^{x}", "2024-25 Half Yearly", 3),
 ("Simplify using the laws of exponents:  (125 × 6^{6} × 10^{4}) / (5^{7} × 9^{3} × 8^{2})", "2022-23 Half Yearly", 3),
 ("P = [ (4/9)^{−2} × (3/4)^{−2} ] × 1/3^{−2} × (17/39)^{0}.\nFind the value of P and also find its reciprocal.", "2022-23 Half Yearly", 3),
 ("Simplify using the laws of exponents and state the laws used:\n [ (1/3^{2})^{−8} × (2/3)^{4} × (81/16)^{−3} ] ÷ 4^{8}", "2022-23 PT-1", 3),
 ("Simplify using the laws of exponents and state the laws used:\n [ 2^{−3} x^{4} y^{3} z^{−1} × 12^{2} x^{2} y^{2} z^{2} ] / [ 9^{2} x^{5} y^{4} z^{−2} ]", "2023-24 Half Yearly", 3),
 ("Simplify using the laws of exponents and find the value of x. Mention the laws used.\n [ (1/2)^{3} × (1/3)^{5} × 9 ] / [ 12^{−1} × 2^{3} × 3^{−6} ] = (2/3)^{x}", "2023-24 Half Yearly", 4),
 ("State the law of exponents used in each, and:\n(a) Simplify: (3^{9})^{2} × 3^{−15}\n(b) Find the value of x^{−12}, if (2/3)^{4} ÷ (2/3)^{6} = (2/3)^{2x}", "2021-22 PT-1", 4),
 ("Evaluate:\n(a) (8/27)^{2/3} ÷ (32)^{−2/5}\n(b) 2^{4} − [ (√4)^{0} × (−2)^{6} ] ÷ 4", "2025-26 Half Yearly", 4),
 ("(i) Subtract 1.73 × 10^{11} − 5.46 × 10^{8} and express the difference in standard form.\n(ii) Find x, if 7^{2x} ÷ 7^{−3} = 7^{23}", "2021-22 Half Yearly", 3),
 ("Simplify and express in scientific notation:  1.987 × 10^{11} − 6.54 × 10^{8}", "2023-24 Half Yearly", 3),
 ("The total land area of India is 3.287263 × 10^{12} m^{2}. If the land area of the southern part of India is 6.3578 × 10^{11} m^{2} and the northern part covers 2.3893 × 10^{12} m^{2}, find the remaining land area in standard form.", "2025-26 Half Yearly", 3),
 ("(a) The Moon and Mars are at a distance of 3.844 × 10^{5} km and 5.45 × 10^{9} m from the Earth respectively. Which is farther from the Earth, and by how much?\n(b) The mass of the Earth is 6.512 × 10^{24} kg and the mass of the Moon is 7.4 × 10^{22} kg approximately. How many times the mass of the Moon is the mass of the Earth?", "2024-25 Half Yearly", 4),
 ("A particular star is at about 810000000 × 10^{8} m away from the Earth. Given that light travels at 3 × 10^{8} m per second:\n(a) Express the distance of the star from the Earth in scientific notation.\n(b) Express the speed of light in usual form.\n(c) How long does light take from that star to reach the Earth? Write this in scientific notation.", "2025-26 Annual", 3),
 ("CASE STUDY — John came up with interesting facts about very large and very small objects. Study the data and answer:\n • A strand of human hair — diameter 0.00005 m\n • A water molecule — volume 2.99 × 10^{−23} cm^{3}\n • Finger nail — rate of growth 9.200 × 10^{−7} mm/sec\n • The Sun — diameter 1,391,000000 m\n(a) John converted the diameter of a strand of human hair into cm and expressed it in scientific notation. Which is correct?\n (i) 5 × 10^{−3} cm (ii) 0.5 × 10^{−5} cm (iii) 5 × 10^{−4} cm (iv) 0.5 × 10^{−4} cm\n(b) The rate of growth of the nail expressed in usual form is\n (i) 9200 mm/sec (ii) 0.000092 mm/sec (iii) 0.00000092 mm/sec (iv) 9200000 mm/sec\n(c) The volume of 10000 water molecules would be\n (i) 2.99 × 10^{−27} cm^{3} (ii) 2.99 × 10^{−28} cm^{3} (iii) 2.99 × 10^{−18} cm^{3} (iv) 2.99 × 10^{−19} cm^{3}\n(d) The diameter of the Sun in scientific notation is\n (i) 1.391 × 10^{9} m (ii) 1.391 × 10^{6} km (iii) 1.391 × 10^{11} cm (iv) all options are true", "2022-23 Half Yearly", 4),
]

ANS1 = [
 "(b) 10404, since 10404 = 102^{2}. (6250, 38943 and 12547 are not perfect squares — note 38943 ends in 3, so it cannot be a square.)",
 "(a) 63.6056. ∛636.056 = 8.6, ∛636056 = 86, ∛0.000636056 = 0.086. Product = 8.6 × 86 × 0.086 = 63.6056",
 "(c) 16. ∛x = 8 ⇒ x = 512, so 8x = 4096 and ∛4096 = 16.",
 "(d) 11 × 3^{3} × 33 = 11 × 3^{3} × 3 × 11 = 11^{2} × 3^{4} = (11 × 3^{2})^{2} = 99^{2}. All prime factors occur in pairs.",
 "(a) 6. 486 = 2 × 3^{5}. Dividing by 2 × 3 = 6 gives 81 = 9^{2}.",
 "(b). If x^{2} ends in 5, then x ends in 5, so x^{3} also ends in 5. (There are only 3 perfect cubes between 1 and 100 — 8, 27, 64; 3^{3} = 27 is a two-digit cube; and the cube root of 8 is 2 only.)",
 "(c) x/100. √11025 = 105, and √1.1025 = 1.05 = 105/100 = x/100.",
 "(c) always a 3-digit number. A 5-digit number lies between 10000 and 99999, so its square root lies between 100 and 316 — always 3 digits. (Options (a) and (b) are identical — a misprint in the original paper.)",
 "(d) 0.027. (0.3)^{3} = 0.3 × 0.3 × 0.3 = 0.027",
 "(c) 8. 25% of 8 = 2, and ∛8 = 2.",
 "(b) 35 and 37. Using 2m = 12 ⇒ m = 6, giving m^{2} − 1 = 35 and m^{2} + 1 = 37. Check: 12^{2} + 35^{2} = 37^{2}.",
 "(c) (ii) and (iii). The number = 2^{6} × 3 × 5^{3}. Multiplying by 3 × 5 gives 2^{6} × 3^{2} × 5^{4}; dividing by 3 × 5 gives 2^{6} × 5^{2}. Both are perfect squares.",
 "9. The unit digit of 3457 is 7, and 7^{2} = 49, so the square ends in 9.",
 "11. √49 = 7, so √(128 − 7) = √121 = 11.",
 "42. 74088 → last group gives unit digit 2; first group 74 lies between 4^{3} = 64 and 5^{3} = 125 ⇒ tens digit 4. So ∛74088 = 42.",
 "281. a^{2} − b^{2} = (a + b)(a − b) = (141 + 140)(141 − 140) = 281 × 1 = 281",
 "21 + 23 + 25 + 27 + 29 = 125 = 5^{3}. (The cube of n is the sum of n consecutive odd numbers.)",
 "21 cm. Height = ∛9261 = 21.",
 "47. (47^{3} = 103823)",
 "23 cm. Side = ∛12167 = 23.",
 "8600.086. √73960000 = 8600 and √0.007396 = 0.086.",
 "24. (24^{3} = 13824)",
 "(a) 1253687 is definitely not a square, because a perfect square can only end in 0, 1, 4, 5, 6 or 9 — never in 2, 3, 7 or 8. This number ends in 7.",
 "Units digit 7, tens digit 8 — ∛658503 = 87.",
 "45. For consecutive numbers, a^{2} − b^{2} = a + b when a − b = 1. So 23^{2} − 22^{2} = 23 + 22 = 45.",
 "(10, 24, 26). Taking m^{2} + 1 = 26 ⇒ m = 5, so 2m = 10 and m^{2} − 1 = 24. Check: 10^{2} + 24^{2} = 26^{2}.",
 "8 and 15. Taking m^{2} + 1 = 17 ⇒ m = 4, so 2m = 8 and m^{2} − 1 = 15. Check: 8^{2} + 15^{2} = 289 = 17^{2}.",
 "16 and 63. Taking m^{2} + 1 = 65 ⇒ m = 8, so 2m = 16 and m^{2} − 1 = 63. Check: 16^{2} + 63^{2} = 65^{2}.",
 "(i) 43  (ii) 27/3^{9} = 3^{3}/3^{9} = 3^{−6}, so the cube root is 3^{−2} = 1/9.",
 "x = 49. 1323 = 3^{3} × 7^{2}. Removing 7^{2} leaves 3^{3} = 27, a perfect cube. So x = 49.",
 "1/2. x = ∛(27/64) = 3/4 and y = ∛(−1/64) = −1/4, so x + y = 3/4 − 1/4 = 1/2.",
 "Multiply by 5; the square root is 210. 8820 = 2^{2} × 3^{2} × 5 × 7^{2}. Only 5 is unpaired, so multiply by 5 → 44100 = 210^{2}.",
 "Subtract 51; the square root is 75. 75^{2} = 5625 and 76^{2} = 5776, so 5676 − 5625 = 51.",
 "770. ∛343 = 7, ∛0.343 = 0.7, ∛10^{6} = 100. So (7 + 0.7) × 100 = 770.",
 "Multiply by 3; the square root is 84. 2352 = 2^{4} × 3 × 7^{2}. Only 3 is unpaired → 2352 × 3 = 7056 = 84^{2}.",
 "25.275. √625 = 25, √0.000625 = 0.025, √0.0625 = 0.25. Sum = 25.275.",
 "3/2 (= 1.5). ∛(18^{3} × (1/12)^{3}) = 18 × 1/12 = 3/2.",
 "27. (27^{3} = 19683)",
 "0.563. √0.3844 = 0.62 and √0.003249 = 0.057, so 0.62 − 0.057 = 0.563.",
 "26. 17576 = 2^{3} × 13^{3}, so ∛17576 = 2 × 13 = 26.",
 "4.47. (4.47^{2} = 19.9809, and 4.48^{2} = 20.0704)",
 "26.3",
 "11.2",
 "45.6",
 "10.4",
 "7.4",
 "90.2",
 "4900 (= 70^{2}). LCM(10, 20, 28) = 140 = 2^{2} × 5 × 7. To make it a square, multiply by 5 × 7 = 35 → 4900.",
 "No, 2560 is not a perfect cube. 2560 = 2^{9} × 5. Multiply by 5^{2} = 25 → 64000 = 40^{3}.",
 "27000 (= 30^{3}). LCM(4, 9, 10) = 180 = 2^{2} × 3^{2} × 5. For a cube each prime needs a power that is a multiple of 3 → 2^{3} × 3^{3} × 5^{3} = 27000.",
 "3600 (= 60^{2}); square root = 60. LCM(80, 12, 24, 36) = 720 = 2^{4} × 3^{2} × 5. Only 5 is unpaired → 720 × 5 = 3600.",
 "Divide by 7; the cube root is 15. 23625 = 3^{3} × 5^{3} × 7. Removing 7 gives 3375 = 15^{3}.",
 "(a) Subtract 4. 86^{2} = 7396, so 7400 − 7396 = 4.\n(b) Add 169. 87^{2} = 7569, so 7569 − 7400 = 169.",
 "(a) 46. Between n^{2} and (n + 1)^{2} there are 2n numbers, so 2 × 23 = 46.\n(b) 5 digits. 161728401 has 9 digits, and the square root of an n-digit number has (n + 1)/2 digits when n is odd.\n(c) 289. The sum of the first n odd numbers is n^{2}, so 17^{2} = 289.",
 "(a) (24, 143, 145), using 2m = 24 ⇒ m = 12. (The well-known triplet (7, 24, 25) also contains 24.)\n(b) 36. 2 × 18 = 36.\n(c) 36 = 1 + 3 + 5 + 7 + 9 + 11 (the first 6 odd numbers, and 6^{2} = 36).",
 "4 cuboids. Volume of one cuboid = 15 × 30 × 15 = 6750 = 2 × 3^{3} × 5^{3}. Only 2 is short of a triple, so multiply by 2^{2} = 4 → 27000 = 30^{3}.",
 "₹4, ₹6 and ₹8. Let the amounts be 2x, 3x, 4x. Then 8x^{3} + 27x^{3} + 64x^{3} = 99x^{3} = 792 ⇒ x^{3} = 8 ⇒ x = 2.",
 "22, 33 and 44. 99x^{3} = 131769 ⇒ x^{3} = 1331 ⇒ x = 11.",
 "154 more persons; 98 persons in each row. 97^{2} = 9409 and 98^{2} = 9604, so 9604 − 9450 = 154.",
 "(a) 9 dancers. 44^{2} = 1936 and 45^{2} = 2025, so 2025 − 2016 = 9.\n(b) 45 dancers in each row.",
 "130 more flowers; 76 flowers in a row. 75^{2} = 5625 and 76^{2} = 5776, so 5776 − 5646 = 130.",
 "6 chairs. 117^{2} = 13689, and 13695 − 13689 = 6.",
 "Perimeter = 336 m; cost = ₹4032. Side = √7056 = 84 m, so perimeter = 4 × 84 = 336 m and cost = 336 × 12 = ₹4032.",
]

ANS2 = [
 "(b) 9p^{−6}. (−3p^{−3})^{2} = (−3)^{2} × (p^{−3})^{2} = 9p^{−6}.",
 "(b) (0.03)^{7}. 0.000027 = (0.03)^{3}, so (0.03)^{−2} × (0.03)^{9} = (0.03)^{7}.",
 "(a) 4^{2}. (2/3)^{2} = 4/9 and (−1/3)^{−2} = 9, so the bracket = 4, and 4^{2} follows.",
 "(c) 10^{3} ÷ 10^{9} = 10^{3−9} = 10^{−6}.",
 "(b) (−7/5)^{−3}. (7/−5)^{3} = (−7/5)^{3}, and a number times its reciprocal is 1, so the blank is (−7/5)^{−3}.",
 "(d) 3. Any power of 1 is 1, so 1 + 1 + 1 = 3.",
 "(a) y/x. x^{−1} ÷ y^{−1} = (1/x) ÷ (1/y) = (1/x) × y = y/x.",
 "(b) 1. 12^{2} − 5^{3} = 144 − 125 = 19, and (−1)^{20} = 1, so 19 × (1/19) = 1.",
 "(a) 1/4. (5/8)^{−2} = 64/25 and (4/5)^{−2} = 25/16, so x = 4 and x^{−1} = 1/4.",
 "(c). 4.26 × 10^{7} = 42600000.",
 "(d) 25. 5^{−2} = 1/25, whose multiplicative inverse is 25.",
 "−1. (−2)^{0} − (−2)^{1} + (−2)^{2} + (−2)^{3} = 1 + 2 + 4 − 8 = −1.",
 "3.412652 × 10^{−5}",
 "3.45 × 10^{7}",
 "5. The blank = (1/5)^{2} × (1/5)^{−3} = (1/5)^{−1} = 5.",
 "3/2. 1^{−1} = 1 and 3^{−1} = 1/3, so (1 − 1/3)^{−1} = (2/3)^{−1} = 3/2.",
 "15260 kg (= 1.526 × 10^{4} kg). 15 × 10^{3} = 15000 and 2.6 × 10^{2} = 260.",
 "a = 6, b = 7. 5b × x^{a+2} = 35x^{8} ⇒ 5b = 35 ⇒ b = 7, and a + 2 = 8 ⇒ a = 6.",
 "4/5. Any non-zero number to the power 0 is 1, so 5^{−1}(1 + 1 + 1) + 5^{−1} = 3/5 + 1/5 = 4/5.",
 "81. (24/3) = 8, so 8^{4} × (3/8)^{4} = (8 × 3/8)^{4} = 3^{4} = 81.",
 "9/64. 2^{−5} = 1/32, 8^{−2} = 1/64, 2^{−3} = 1/8. So 2/64 − 1/64 + 8/64 = 9/64.",
 "1. (−2/3)^{3} × (−2/3) = (−2/3)^{4} = 16/81, and (4/9)^{2} = 16/81. So the quotient is 1.",
 "1/5. (5^{2})^{−1} × 125 = 5^{−2} × 5^{3} = 5, and (1/5)^{−2} = 25. So 5 ÷ 25 = 1/5.",
 "p^{4} q^{4} r^{4}. Subtract exponents: p^{7−3} q^{2−(−2)} r^{−4−(−8)} = p^{4} q^{4} r^{4}.",
 "x = 7. LHS = 5^{−8+3x}, RHS = 5^{6+x}. So −8 + 3x = 6 + x ⇒ 2x = 14 ⇒ x = 7.",
 "p = 8. (1/12)^{−5} = 12^{5}, so 12^{p−3} = 12^{5} ⇒ p = 8.",
 "x = 2. (5/4)^{3} = (4/5)^{−3}, so LHS = (4/5)^{3−(−3)} = (4/5)^{6}. Hence 3x = 6 ⇒ x = 2.",
 "x = 1. (2/5)^{−2} = 25/4, and 625/16 = (25/4)^{2}. So 2x = 2 ⇒ x = 1.",
 "64/729. a/b = (3/2)^{3} × 1 = 27/8, so (a/b)^{−2} = (8/27)^{2} = 64/729.",
 "x = 12. (81/25)^{−3} = (25/81)^{3} = (5/9)^{6}, and LHS = (5/9)^{2x−18}. So 2x − 18 = 6 ⇒ x = 12.",
 "x = −1. (9/16)^{4} = (3/4)^{8} = (4/3)^{−8} and 64/27 = (4/3)^{3}. So LHS = (4/3)^{−8+10−3} = (4/3)^{−1}.",
 "16. Numerator = 5^{3} × 2^{6}3^{6} × 2^{4}5^{4} = 2^{10} 3^{6} 5^{7}; denominator = 5^{7} × 3^{6} × 2^{6}. So the result is 2^{4} = 16.",
 "P = 81; reciprocal = 1/81. (4/9)^{−2} = 81/16 and (3/4)^{−2} = 16/9, whose product is 9. Also 1/3^{−2} = 9 and (17/39)^{0} = 1. So P = 9 × 9 = 81.",
 "1. (1/3^{2})^{−8} = 3^{16}; (2/3)^{4} = 2^{4}3^{−4}; (81/16)^{−3} = 2^{12}3^{−12}; 4^{8} = 2^{16}. Numerator = 2^{16} 3^{0} = 2^{16}, so the quotient is 1.\nLaws used: (a^{m})^{n} = a^{mn}, a^{m} × a^{n} = a^{m+n}, a^{−n} = 1/a^{n}, a^{m} ÷ a^{n} = a^{m−n}.",
 "(2/9) x y z^{3}. Coefficients: 2^{−3} × 144 ÷ 81 = 18/81 = 2/9. Then x^{4+2−5} = x, y^{3+2−4} = y, z^{−1+2+2} = z^{3}.\nLaws used: a^{m} × a^{n} = a^{m+n} and a^{m} ÷ a^{n} = a^{m−n}.",
 "x = −4. Numerator = 2^{−3} × 3^{−5} × 3^{2} = 2^{−3}3^{−3}. Denominator = 2^{−2}3^{−1} × 2^{3} × 3^{−6} = 2^{1}3^{−7}. So LHS = 2^{−4}3^{4} = (2/3)^{−4}.",
 "(a) 27. (3^{9})^{2} × 3^{−15} = 3^{18−15} = 3^{3} = 27. Laws: (a^{m})^{n} = a^{mn} and a^{m} × a^{n} = a^{m+n}.\n(b) 1. (2/3)^{4−6} = (2/3)^{−2}, so 2x = −2 ⇒ x = −1, and x^{−12} = (−1)^{−12} = 1.",
 "(a) 16/9. (8/27)^{2/3} = (2/3)^{2} = 4/9, and 32^{−2/5} = 2^{−2} = 1/4. So (4/9) ÷ (1/4) = 16/9.\n(b) 0. 2^{4} = 16; (√4)^{0} = 1 and (−2)^{6} = 64, so [1 × 64] ÷ 4 = 16. Hence 16 − 16 = 0.",
 "(i) 1.72454 × 10^{11}. 1.73 × 10^{11} − 5.46 × 10^{8} = (1730 − 5.46) × 10^{8} = 1724.54 × 10^{8}.\n(ii) x = 10. 2x − (−3) = 23 ⇒ 2x = 20 ⇒ x = 10.",
 "1.98046 × 10^{11}. (1987 − 6.54) × 10^{8} = 1980.46 × 10^{8}.",
 "2.62183 × 10^{11} m^{2}. Southern + northern = 0.63578 × 10^{12} + 2.3893 × 10^{12} = 3.02508 × 10^{12}. Remaining = 3.287263 × 10^{12} − 3.02508 × 10^{12} = 0.262183 × 10^{12}.",
 "(a) Mars is farther, by 5.0656 × 10^{9} m. The Moon is 3.844 × 10^{5} km = 3.844 × 10^{8} m, so 5.45 × 10^{9} − 0.3844 × 10^{9} = 5.0656 × 10^{9} m.\n(b) About 88 times. (6.512 × 10^{24}) ÷ (7.4 × 10^{22}) = 0.88 × 10^{2} = 88.",
 "(a) 8.1 × 10^{16} m. 810000000 × 10^{8} = 8.1 × 10^{8} × 10^{8}.\n(b) 300000000 m/s.\n(c) 2.7 × 10^{8} seconds. Time = distance ÷ speed = (8.1 × 10^{16}) ÷ (3 × 10^{8}).",
 "(a) (i) 5 × 10^{−3} cm. 0.00005 m = 0.005 cm.\n(b) (iii) 0.00000092 mm/sec.\n(c) (iv) 2.99 × 10^{−19} cm^{3}. 2.99 × 10^{−23} × 10^{4} = 2.99 × 10^{−19}.\n(d) (iv) all options are true. 1,391,000000 m = 1.391 × 10^{9} m = 1.391 × 10^{6} km = 1.391 × 10^{11} cm.",
]

# ---------------------------------------------------------------- rendering
TOKEN = re.compile(r"\^\{([^}]*)\}|_\{([^}]*)\}")


def rich(par, text, size=10.5, bold=False, italic=False, color=None):
    """Write text into a paragraph, honouring ^{sup} and _{sub} markers."""
    pos = 0
    for m in TOKEN.finditer(text):
        if m.start() > pos:
            r = par.add_run(text[pos:m.start()])
            r.font.size, r.bold, r.italic = Pt(size), bold, italic
            if color: r.font.color.rgb = color
        r = par.add_run(m.group(1) or m.group(2))
        r.font.size, r.bold, r.italic = Pt(size), bold, italic
        if m.group(1): r.font.superscript = True
        else: r.font.subscript = True
        if color: r.font.color.rgb = color
        pos = m.end()
    if pos < len(text):
        r = par.add_run(text[pos:])
        r.font.size, r.bold, r.italic = Pt(size), bold, italic
        if color: r.font.color.rgb = color


NAVY = RGBColor(0x1F, 0x3A, 0x5F)
GREY = RGBColor(0x70, 0x70, 0x70)
GREEN = RGBColor(0x1B, 0x5E, 0x20)

doc = Document()
for s in doc.sections:
    s.top_margin = s.bottom_margin = Inches(0.7)
    s.left_margin = s.right_margin = Inches(0.8)
doc.styles["Normal"].font.name = "Calibri"
doc.styles["Normal"].font.size = Pt(10.5)

# title
t = doc.add_paragraph(); t.alignment = WD_ALIGN_PARAGRAPH.CENTER
rich(t, "Class 8 Mathematics — Question Bank", size=20, bold=True, color=NAVY)
s = doc.add_paragraph(); s.alignment = WD_ALIGN_PARAGRAPH.CENTER
rich(s, "Chapter 1: A Square and A Cube • Chapter 2: Power Play", size=13, color=NAVY)
s2 = doc.add_paragraph(); s2.alignment = WD_ALIGN_PARAGRAPH.CENTER
rich(s2, "Compiled from Sri Kumaran Children's Home (CBSE) past papers, 2021-22 to 2025-26", size=9.5, italic=True, color=GREY)
n = doc.add_paragraph(); n.alignment = WD_ALIGN_PARAGRAPH.CENTER
rich(n, f"{len(CH1)} questions on Squares & Cubes • {len(CH2)} questions on Exponents & Powers • Answers at the end", size=9.5, color=GREY)
doc.add_paragraph()
lh = doc.add_paragraph(); lh.alignment = WD_ALIGN_PARAGRAPH.CENTER
rich(lh, "How to read the stars", size=11.5, bold=True, color=NAVY)
for _sym, _txt, _col in [
    ("***", "MUST KNOW - this skill came up in 7 or more of the exams. If time is short, revise these first.", RGBColor(0xC6, 0x28, 0x28)),
    ("**", "IMPORTANT - came up in 4 to 6 of the exams.", RGBColor(0xB2, 0x5E, 0x00)),
    ("*", "SEEN SOMETIMES - came up in 2 or 3 of the exams.", GREY),
    ("(no star)", "Appeared just once. Do these last.", GREY),
]:
    _p = doc.add_paragraph(); _p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _p.paragraph_format.space_after = Pt(2)
    rich(_p, f"{_sym}  ", size=10, bold=True, color=_col)
    rich(_p, _txt, size=9.5, color=GREY)
_p = doc.add_paragraph(); _p.alignment = WD_ALIGN_PARAGRAPH.CENTER
rich(_p, "The stars are not an opinion - they were counted from the papers themselves.", size=9, italic=True, color=GREY)


from topics import T1, T2, analyse, stars, TIER_NAME
from docx.enum.table import WD_TABLE_ALIGNMENT

RED = RGBColor(0xC6, 0x28, 0x28)
AMBER = RGBColor(0xB2, 0x5E, 0x00)
TIER_COLOR = {"***": RED, "**": AMBER, "*": GREY, "": GREY}


def priority_table(title, items, tags):
    h = doc.add_paragraph()
    rich(h, title, size=12.5, bold=True, color=NAVY)
    n_exams = len({s for _, s, _ in items})
    c = doc.add_paragraph()
    rich(c, f"Worked out by counting how often each skill actually appeared across the {n_exams} different exams in this bank.",
         size=9, italic=True, color=GREY)
    tb = doc.add_table(rows=1, cols=4)
    tb.style = "Light Grid Accent 1"
    tb.alignment = WD_TABLE_ALIGNMENT.CENTER
    for c_, htxt in enumerate(["Priority", "Topic", "Questions", "Exams it appeared in"]):
        cell = tb.rows[0].cells[c_]
        cell.paragraphs[0].text = ""
        rich(cell.paragraphs[0], htxt, size=9, bold=True)
    for t, cnt, ex, _ in analyse(items, tags):
        st = stars(ex)
        cells = tb.add_row().cells
        for c_, v in enumerate([f"{st}  {TIER_NAME[st]}".strip(), t, str(cnt), f"{ex} of {n_exams}"]):
            cells[c_].paragraphs[0].text = ""
            rich(cells[c_].paragraphs[0], v, size=9,
                 bold=(c_ == 0 and st == "***"), color=TIER_COLOR[st] if c_ == 0 else None)
    doc.add_paragraph()


def section(title, subtitle, items, tags, start):
    doc.add_page_break()
    h = doc.add_paragraph(); rich(h, title, size=16, bold=True, color=NAVY)
    sh = doc.add_paragraph(); rich(sh, subtitle, size=9.5, italic=True, color=GREY)
    doc.add_paragraph()
    priority_table("Where the marks are — start with the ***", items, tags)

    exam_count = {t: ex for t, _, ex, _ in analyse(items, tags)}
    doc.add_page_break()
    for i, ((q, src, marks), tag) in enumerate(zip(items, tags), start):
        st = stars(exam_count[tag])
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        rich(p, f"{i}. ", size=10.5, bold=True)
        if st:
            rich(p, f"{st} ", size=10.5, bold=True, color=TIER_COLOR[st])
        rich(p, q, size=10.5)
        m = doc.add_paragraph()
        m.paragraph_format.space_after = Pt(9)
        rich(m, f"[{src} — {marks} mark{'s' if marks > 1 else ''}]  ", size=8.5, italic=True, color=GREY)
        rich(m, tag, size=8.5, italic=True, color=TIER_COLOR[st])


section("Chapter 1 — A Square and A Cube",
        "Squares • square roots • cubes • cube roots • Pythagorean triplets", CH1, T1, 1)
section("Chapter 2 — Power Play",
        "Exponents • laws of exponents • negative powers • standard form / scientific notation", CH2, T2, 1)

# answers
doc.add_page_break()
h = doc.add_paragraph(); rich(h, "Answers", size=16, bold=True, color=GREEN)
sh = doc.add_paragraph()
rich(sh, "Try each question fully before checking. Working is shown where it helps.", size=9.5, italic=True, color=GREY)

for name, ans in (("Chapter 1 — A Square and A Cube", ANS1), ("Chapter 2 — Power Play", ANS2)):
    doc.add_paragraph()
    ph = doc.add_paragraph(); rich(ph, name, size=12.5, bold=True, color=NAVY)
    for i, a in enumerate(ans, 1):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(5)
        rich(p, f"{i}. ", size=10, bold=True)
        rich(p, a, size=10)

assert len(CH1) == len(ANS1), "Ch1 mismatch!"
assert len(CH2) == len(ANS2), "Ch2 mismatch!"

if __name__ == "__main__":
    doc.save(OUT)
    print("saved:", OUT)
    print(f"Ch1: {len(CH1)} questions / {len(ANS1)} answers")
    print(f"Ch2: {len(CH2)} questions / {len(ANS2)} answers")
    print("counts match ✓")
