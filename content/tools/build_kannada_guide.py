#!/usr/bin/env python3
"""A tough-test Kannada study guide (bilingual: English explanation, Kannada term),
matching the Science guide. Marked by how often each topic appears in the 49 real
Sri Kumaran Kannada papers.

Every Kannada character is pulled from the verified JSON (content/data/
kannada_source.json and kannada_worksheets.json) — NOT hand-typed. Kannada was
OCR'd and human-checked earlier; hand-typing it introduces errors.

Kannada renders in 'Kannada Sangam MN' (a macOS system font); the complex-script
font attribute (w:cs) is set so Word uses it for the Kannada glyphs.
"""
import json, os, re
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

HERE = os.path.dirname(os.path.abspath(__file__))
SRC = json.load(open(os.path.join(HERE, "..", "data", "kannada_source.json")))
WS = json.load(open(os.path.join(HERE, "..", "data", "kannada_worksheets.json")))
OUT = "/Users/shyamdk/Developer/personal/laya/data/question-papers/Kannada - Study Guide (Lessons + Grammar).docx"

KAN_FONT = "Kannada Sangam MN"
KAN = re.compile(r"[ಀ-೿]")

NAVY = RGBColor(0x1F, 0x3A, 0x5F)
GREY = RGBColor(0x60, 0x60, 0x60)
GREEN = RGBColor(0x1B, 0x5E, 0x20)
RED = RGBColor(0xC6, 0x28, 0x28)
AMBER = RGBColor(0xB2, 0x5E, 0x00)
TEAL = RGBColor(0x00, 0x69, 0x5C)

doc = Document()
for s in doc.sections:
    s.top_margin = s.bottom_margin = Inches(0.7)
    s.left_margin = s.right_margin = Inches(0.8)
doc.styles["Normal"].font.name = "Calibri"
doc.styles["Normal"].font.size = Pt(10.5)


def _style_run(r, size, bold, italic, color, kannada):
    r.font.size = Pt(size)
    r.bold = bold
    r.italic = italic
    if color:
        r.font.color.rgb = color
    if kannada:
        r.font.name = KAN_FONT
        # Kannada is a complex script — Word picks the font from w:cs, so set it.
        rpr = r._element.get_or_add_rPr()
        rfonts = rpr.find(qn("w:rFonts"))
        if rfonts is None:
            rfonts = rpr.makeelement(qn("w:rFonts"), {})
            rpr.append(rfonts)
        for attr in ("w:ascii", "w:hAnsi", "w:cs"):
            rfonts.set(qn(attr), KAN_FONT)


def write(p, text, size=10.5, bold=False, italic=False, color=None):
    """Write text into paragraph p, switching to the Kannada font per script run."""
    if not text:
        return
    # split into maximal Kannada / non-Kannada spans (keep spaces with either)
    parts = re.split(r"((?:[ಀ-೿][ಀ-೿\s,।]*)+)", text)
    for part in parts:
        if not part:
            continue
        _style_run(p.add_run(part), size, bold, italic, color, bool(KAN.search(part)))


def para(text="", **kw):
    after = kw.pop("after", 4)
    align = kw.pop("align", None)
    indent = kw.pop("indent", None)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(after)
    if align:
        p.alignment = align
    if indent:
        p.paragraph_format.left_indent = Inches(indent)
    write(p, text, **kw)
    return p


def h1(text, accent=NAVY):
    doc.add_page_break()
    para(text, size=19, bold=True, color=accent, after=3)


def h2(text, tier=None, why=None, accent=NAVY):
    para(text, size=13.5, bold=True, color=accent, after=2)
    if tier:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(4)
        col, lab = {"***": (RED, "MUST KNOW"), "**": (AMBER, "IMPORTANT"),
                    "*": (GREY, "SEEN SOMETIMES")}[tier]
        r = p.add_run(f"{tier} {lab}  ")
        r.bold, r.font.size, r.font.color.rgb = True, Pt(9), col
        if why:
            write(p, why, size=9, italic=True, color=GREY)


def box(label, text, color):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.12)
    p.paragraph_format.space_after = Pt(5)
    r = p.add_run(f"{label}  ")
    r.bold, r.font.size, r.font.color.rgb = True, Pt(9), color
    write(p, text, size=10, color=color)


def bullets(items, color=None):
    for it in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(2)
        write(p, it, size=10, color=color)


def table(head, rows):
    tb = doc.add_table(rows=1, cols=len(head))
    tb.style = "Light Grid Accent 1"
    tb.alignment = WD_TABLE_ALIGNMENT.CENTER
    for c, h in enumerate(head):
        tb.rows[0].cells[c].paragraphs[0].text = ""
        write(tb.rows[0].cells[c].paragraphs[0], h, size=9.5, bold=True)
    for row in rows:
        cells = tb.add_row().cells
        for c, v in enumerate(row):
            cells[c].paragraphs[0].text = ""
            write(cells[c].paragraphs[0], v, size=10)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)


L = {l["code"]: l for l in SRC["lessons"]}
G = {g["code"]: g for g in SRC["grammar_topics"]}
FREQ = SRC["exam_frequency_counts"]

# ============================================================ COVER
para("ಕನ್ನಡ — Study Guide", size=24, bold=True, color=NAVY, align=WD_ALIGN_PARAGRAPH.CENTER, after=2)
para("Grade 8 · Second Language · for the coming test", size=12, color=NAVY,
     align=WD_ALIGN_PARAGRAPH.CENTER, after=2)
para("ಮಗ್ಗದ ಸಾಹೇಬ  ·  ಕನ್ನಡಿಗರ ತಾಯಿ  ·  ಸಂಧಿ · ಸಮಾಸ · ತತ್ಸಮ-ತದ್ಭವ",
     size=11, italic=True, color=GREY, align=WD_ALIGN_PARAGRAPH.CENTER, after=10)

para("What to study first — counted from 49 real Kannada papers (2009–2026)",
     size=12, bold=True, color=NAVY)
para("The stars are not a guess. They are how many of the 49 past papers each topic appeared in. "
     "Study the ✦✦✦ topics first.", size=9.5, italic=True, color=GREY, after=6)
table(["Topic", "In the papers", "Priority"], [
    ["ಸಂಧಿ  (joining words)", f"{FREQ['kn-g-sandhi']['papers']} of 49", "*** MUST KNOW"],
    ["ಪದಗಳ ಅರ್ಥ  (word meanings, both lessons)", f"{FREQ['vocabulary']['papers']} of 49", "*** MUST KNOW"],
    ["ತತ್ಸಮ – ತದ್ಭವ  (Sanskrit-borrowed words)", f"{FREQ['kn-g-tatsama']['papers']} of 49", "*** MUST KNOW"],
    ["ಸಮಾಸ  (compound words)", f"{FREQ['kn-g-samasa']['papers']} of 49", "** IMPORTANT"],
    ["Lesson comprehension  (the two ಪಾಠಗಳು)", "school worksheets", "*** MUST KNOW"],
    ["ವರ್ಣಮಾಲೆ / ಗುಣಿತಾಕ್ಷರ  (alphabet)", "0 of 49", "* SEEN SOMETIMES"],
])
box("HOW TO USE THIS", "Read the coloured boxes: KEY FACT (green) is what you must be able to write, "
    "TRAP (amber) is where marks are lost, and the RULE boxes tell you HOW to identify each grammar "
    "type. English explains; the Kannada is what you write in the exam. Each part ends with a self-test.",
    NAVY)

# ============================================================ helper for lesson comprehension
def lesson_facts(code):
    return [k for k in WS["kept"] if k["lesson"] == code]

# ============================================================ LESSON 1
l1 = L["kn-l1"]
h1("1.  ಗದ್ಯ — ಮಗ್ಗದ ಸಾಹೇಬ", accent=GREEN)
para(f"Author (ಲೇಖಕರು): {l1['author_kn']}.  A prose lesson (ಗದ್ಯ).", size=10, italic=True, color=GREY, after=6)

h2("1.1  ಪದಗಳ ಅರ್ಥ — Word meanings", "***",
   "Word meanings appeared in 33 of the 49 papers — and the wrong options are usually the OTHER words in this list.", GREEN)
para("Learn these cold. In the paper you are given a word and must pick its meaning (or the reverse).")
table(["ಪದ (word)", "ಅರ್ಥ (Kannada meaning)", "English"],
      [[w["word"], w["meaning"], w.get("english", "")] for w in l1["word_meanings"]])

h2("1.2  ಟಿಪ್ಪಣಿ — Notes (named terms)", "**", "Short 'explain this term' questions come from here.", GREEN)
for n in l1["notes"]:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    write(p, n["term"] + " : ", size=10, bold=True, color=NAVY)
    write(p, n["note"], size=10)

h2("1.3  Key comprehension points (from the school's own worksheet)", "***",
   "These are from the practice sheet the school gave — the closest thing to the real test.", GREEN)
para("The lesson is about Karim, whose father Abdul Rahim hated the loom. Know these facts:")
facts1 = lesson_facts("kn-l1")
table(["Question", "Answer"], [[f["stem"], f["options"][f["answer_index"]]] for f in facts1])

para("Rapid revision — ಮಗ್ಗದ ಸಾಹೇಬ", size=11, bold=True, color=GREEN, after=2)
bullets([
    f"Author: {l1['author_kn']}, from ಬಾಗಲೋಡಿ village (his surname is his village).",
    "Theme: a father hates the loom; his son Karim learns it, masters it, and brings the family honour.",
    "Karim's headmaster: ಶಂಕರಪ್ಪ.  The 'new education' was inspired by ಮಹಾತ್ಮ ಗಾಂಧಿ.",
    "Karim won a silver medal (ಬೆಳ್ಳಿಯ ಪದಕ) and ₹100; later the Padma Bhushan (ಪದ್ಮಭೂಷಣ).",
    "ವಿಲಾಯತಿ = ವಿದೇಶ.  ಕಳೇಬರ = ಹೆಣ / ಮೃತ ದೇಹ.  ಹಸನ್ಮುಖಿ = ನಗುಮುಖದವನು.",
])

para("Self-test — ಮಗ್ಗದ ಸಾಹೇಬ  (answers at the end)", size=11, bold=True, color=GREEN, after=3)
for i, q in enumerate([
    "ಮಗ್ಗದ ಸಾಹೇಬ ಪಾಠದ ಲೇಖಕರು ಯಾರು?  (Who wrote this lesson?)",
    "'ವಿಲಾಯತಿ' ಪದದ ಅರ್ಥ ಬರೆಯಿರಿ.",
    "ಕರೀಮನಿಗೆ ರಾಷ್ಟ್ರಪತಿಗಳು ನೀಡಿದ ಅತ್ಯುನ್ನತ ಪ್ರಶಸ್ತಿ ಯಾವುದು?",
    "'ಕಳೇಬರ' ಪದದ ಅರ್ಥವೇನು?",
], 1):
    para(f"{i}. {q}", size=10, indent=0.12, after=2)

# ============================================================ LESSON 5
l5 = L["kn-l5"]
h1("2.  ಪದ್ಯ — ಕನ್ನಡಿಗರ ತಾಯಿ", accent=GREEN)
para(f"Poet (ಕವಿ): {l5['author_kn']}.  A poem (ಪದ್ಯ) praising the Kannada mother-figure.",
     size=10, italic=True, color=GREY, after=6)

h2("2.1  ಪದಗಳ ಅರ್ಥ — Word meanings", "***", "33 of 49 papers.", GREEN)
table(["ಪದ (word)", "ಅರ್ಥ (Kannada meaning)", "English"],
      [[w["word"], w["meaning"], w.get("english", "")] for w in l5["word_meanings"]])

h2("2.2  ಟಿಪ್ಪಣಿ — Notes (compound-word meanings)", "**", None, GREEN)
for n in l5["notes"]:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    write(p, n["term"] + " : ", size=10, bold=True, color=NAVY)
    write(p, n["note"], size=10)

facts5 = lesson_facts("kn-l5")
if facts5:
    h2("2.3  Key comprehension points", "***", "From the school worksheet.", GREEN)
    table(["Question", "Answer"], [[f["stem"], f["options"][f["answer_index"]]] for f in facts5])

para("Rapid revision — ಕನ್ನಡಿಗರ ತಾಯಿ", size=11, bold=True, color=GREEN, after=2)
bullets([
    f"Poet: {l5['author_kn']} — the ರಾಷ್ಟ್ರಕವಿ (national poet), Kannada's FIRST Rashtrakavi.",
    "Theme: the Kannada mother should show her face to the world; Kannada is the Kannadigas' very breath.",
    "ಬೇಲನಾಡು = ಬೇಲೂರು.  ಶರ್ವ = another name for ನೃಪತುಂಗ.  ನಂದಿನಿ = ಗೋವು.",
    "ಬಿಳಿಯ ಕೊಳ = ಶ್ರವಣಬೆಳಗೊಳ (white 'pond' — the place, not 'rot').",
])

para("Self-test — ಕನ್ನಡಿಗರ ತಾಯಿ", size=11, bold=True, color=GREEN, after=3)
for i, q in enumerate([
    "ಕನ್ನಡಿಗರ ತಾಯಿ ಪದ್ಯದ ಕವಿ ಯಾರು?",
    "'ಲತೆ' ಪದದ ಅರ್ಥ ಬರೆಯಿರಿ.",
    "'ನಂದಿನಿ' ಎಂದರೆ ಏನು?",
], 1):
    para(f"{i}. {q}", size=10, indent=0.12, after=2)

# ============================================================ GRAMMAR: helper
def grammar_section(code, num, tier):
    g = G[code]
    freq = FREQ.get(code, {})
    why = (f"Appeared in {freq['papers']} of the 49 papers." if freq.get("papers")
           else "Taught with the lesson; rarely asked directly.")
    h1(f"{num}.  ವ್ಯಾಕರಣ — {g['name_kn']}  ({g['name_en'].split(' - ')[0]})", accent=NAVY)
    h2(g["name_en"], tier, why, NAVY)
    para(g["explain_en"])
    for t in g.get("types", []):
        h2(f"{t['name_kn']} — {t['name_en']}", accent=NAVY)
        if t.get("rule_en"):
            box("RULE", t["rule_en"], TEAL)
        rows = [[e["split"], e["joined"]] for e in t["examples"] if "stays apart" not in e["joined"]]
        if rows:
            table(["ಬಿಡಿಸಿ (split)", "ಜೋಡಿಸಿ (joined)"], rows)
    para(f"From the textbook: {g['taught_in']}.", size=9, italic=True, color=GREY, after=4)
    return g


# ============================================================ SANDHI (top priority)
g = grammar_section("kn-g-sandhi", 3, "***")
box("KEY FACT", "SANDHI joins two words into one. Three things can happen at the join, and the sandhi is "
    "named after it:  a letter is DROPPED (ಲೋಪ),  a letter is ADDED (ಆಗಮ: ಯ್ or ವ್),  or a letter is "
    "REPLACED (ಆದೇಶ: ಕ ತ ಪ → ಗ ದ ಬ).", GREEN)
box("TRAP", "In ಆಗಮಸಂಧಿ the ADDED letter is ಯ್ or ವ್.  In ಆದೇಶಸಂಧಿ a consonant is SWAPPED (ಕ→ಗ, ತ→ದ, "
    "ಪ→ಬ). ಪ್ರಕೃತಿಭಾವ means NO sandhi happens at all — the words stay apart.", AMBER)
para("How to decide the type (a reliable method):", size=10, bold=True, after=2)
bullets([
    "Did a letter DISAPPEAR at the join? → ಲೋಪಸಂಧಿ  (ಊರು+ಇಂದ → ಊರಿಂದ).",
    "Did a NEW ಯ್/ವ್ appear? → ಆಗಮಸಂಧಿ  (ಮನೆ+ಅಲ್ಲಿ → ಮನೆಯಲ್ಲಿ).",
    "Did a consonant CHANGE (ಕ→ಗ etc.)? → ಆದೇಶಸಂಧಿ  (ಮಳೆ+ಕಾಲ → ಮಳೆಗಾಲ).",
])
para("Self-test — ಸಂಧಿ", size=11, bold=True, color=NAVY, after=3)
for i, q in enumerate([
    "ಮನೆ + ಅಲ್ಲಿ ಸೇರಿಸಿ, ಸಂಧಿಯ ಹೆಸರು ಬರೆಯಿರಿ.",
    "ಮಳೆ + ಕಾಲ = ?  ಇದು ಯಾವ ಸಂಧಿ?",
    "'ಊರಿಂದ' ಪದವನ್ನು ಬಿಡಿಸಿ ಬರೆಯಿರಿ. ಯಾವ ಸಂಧಿ?",
], 1):
    para(f"{i}. {q}", size=10, indent=0.12, after=2)

# ============================================================ SAMASA
g = grammar_section("kn-g-samasa", 4, "**")
box("KEY FACT", "SAMASA fuses two words into one, dropping the case-ending of the first. Name it by the "
    "FIRST word: describes the second → ಕರ್ಮಧಾರಯ; is a NUMBER → ದ್ವಿಗು; otherwise the SECOND word "
    "carries the meaning → ತತ್ಪುರುಷ.", GREEN)
box("TRAP", "Splitting a samasa back into its parts is the ವಿಗ್ರಹವಾಕ್ಯ. Mixing a Kannada word with a "
    "Sanskrit word is a fault called ಅರಿಸಮಾಸ.", AMBER)
para("Self-test — ಸಮಾಸ", size=11, bold=True, color=NAVY, after=3)
for i, q in enumerate([
    "ಮೂರು + ಮಡಿ = ಮುಮ್ಮಡಿ — ಇದು ಯಾವ ಸಮಾಸ?  (hint: first word is a number)",
    "ಹಿರಿದು + ಮರ = ? ಯಾವ ಸಮಾಸ?",
], 1):
    para(f"{i}. {q}", size=10, indent=0.12, after=2)

# ============================================================ TATSAMA-TADBHAVA
g = grammar_section("kn-g-tatsama", 5, "***")
box("KEY FACT", "TATSAMA = a Sanskrit word taken into Kannada UNCHANGED (ರಾಮ, ಚಂದ್ರ). TADBHAVA = a "
    "Sanskrit word that CHANGED on the way in (ಸಿರಿ, ಬಸದಿ). ANYADESHYA = a word from another language "
    "entirely (English/Hindustani/Portuguese: ಬ್ಯಾಂಕು, ಕಚೇರಿ, ಮೇಜು).", GREEN)
box("TRAP", "A common question gives a word and asks for its ತದ್ಭವ (or ತತ್ಸಮ) form. Learn the pairs — "
    "e.g. the ತತ್ಸಮ 'ಸ್ತ್ರೀ' vs everyday forms; ತದ್ಭವ words like ಸಿರಿ, ಬಾವಿ, ದನಿ change from Sanskrit.",
    AMBER)
para("Self-test — ತತ್ಸಮ-ತದ್ಭವ", size=11, bold=True, color=NAVY, after=3)
for i, q in enumerate([
    "ಈ ಪದಗಳು ತತ್ಸಮವೋ ತದ್ಭವವೋ? — ರಾಮ, ಸಿರಿ, ಚಂದ್ರ, ಬಸದಿ.",
    "'ಬ್ಯಾಂಕು' ಪದ ಯಾವ ವರ್ಗಕ್ಕೆ ಸೇರುತ್ತದೆ?  (Sanskrit-changed, or from another language?)",
], 1):
    para(f"{i}. {q}", size=10, indent=0.12, after=2)

# ============================================================ VARNAMALE / GUNITAKSHARA (quick)
h1("6.  ವ್ಯಾಕರಣ — ವರ್ಣಮಾಲೆ & ಗುಣಿತಾಕ್ಷರ  (quick)", accent=NAVY)
para("Rarely asked directly (0 of 49 papers), but easy marks if it comes.", size=10, italic=True, color=GREY, after=6)
for code in ("kn-g-varnamale", "kn-g-gunitakshara"):
    gg = G[code]
    h2(f"{gg['name_kn']} — {gg['name_en']}", "*", None, NAVY)
    para(gg["explain_en"])

# ============================================================ ANSWERS
h1("Self-test answers", accent=GREEN)
groups = [
    ("ಮಗ್ಗದ ಸಾಹೇಬ", ["ಬಾಗಲೋಡಿ ದೇವರಾಯ.", "ವಿಲಾಯತಿ = ವಿದೇಶ / ಅನ್ಯದೇಶ.",
                      "ಪದ್ಮಭೂಷಣ.", "ಕಳೇಬರ = ಹೆಣ / ಮೃತ ದೇಹ / ಶರೀರ."]),
    ("ಕನ್ನಡಿಗರ ತಾಯಿ", ["ರಾಷ್ಟ್ರಕವಿ ಎಂ. ಗೋವಿಂದ ಪೈ.", "ಲತೆ = ಬಳ್ಳಿ.", "ನಂದಿನಿ = ಗೋವು."]),
    ("ಸಂಧಿ", ["ಮನೆ + ಅಲ್ಲಿ = ಮನೆಯಲ್ಲಿ — ಆಗಮಸಂಧಿ (ಯ್ added).",
              "ಮಳೆ + ಕಾಲ = ಮಳೆಗಾಲ — ಆದೇಶಸಂಧಿ (ಕ → ಗ).",
              "ಊರಿಂದ = ಊರು + ಇಂದ — ಲೋಪಸಂಧಿ (ಉ dropped)."]),
    ("ಸಮಾಸ", ["ದ್ವಿಗು ಸಮಾಸ (first word is a number).",
              "ಹೆಮ್ಮರ — ಕರ್ಮಧಾರಯ ಸಮಾಸ (first word describes the second)."]),
    ("ತತ್ಸಮ-ತದ್ಭವ", ["ತತ್ಸಮ: ರಾಮ, ಚಂದ್ರ.  ತದ್ಭವ: ಸಿರಿ, ಬಸದಿ.",
                     "ಅನ್ಯದೇಶ್ಯ — from English."]),
]
for title, answers in groups:
    para(title, size=12, bold=True, color=GREEN, after=2)
    for a in answers:
        para("• " + a, size=10, color=GREEN, indent=0.1, after=2)

doc.save(OUT)
print("saved:", OUT)
print(f"  {len(doc.paragraphs)} paragraphs, {len(doc.tables)} tables")
