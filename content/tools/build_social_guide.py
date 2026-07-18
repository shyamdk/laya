#!/usr/bin/env python3
"""A tough-test Social Science study guide, marked by exam frequency.

Frequency is counted from the 51 real Sri Kumaran Social papers WHERE MEANINGFUL:
  Geography — natural resources  32/51  *** MUST KNOW
  Civics    — Lok/Rajya Sabha    17/51  ** IMPORTANT (elections)
  Geography — minerals           18/51  ** IMPORTANT
History (Ch2) is NEW to the 2026-27 syllabus — old papers never tested it (0/51),
so it is marked by textbook emphasis and flagged as new, not by a zero frequency.
"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

OUT = "/Users/shyamdk/Developer/personal/laya/data/question-papers/Social Science - Study Guide (Resources, History, Civics).docx"

NAVY = RGBColor(0x1F, 0x3A, 0x5F)
GREY = RGBColor(0x60, 0x60, 0x60)
GREEN = RGBColor(0x1B, 0x5E, 0x20)
RED = RGBColor(0xC6, 0x28, 0x28)
AMBER = RGBColor(0xB2, 0x5E, 0x00)
TEAL = RGBColor(0x00, 0x69, 0x5C)
BROWN = RGBColor(0x8D, 0x4E, 0x00)

doc = Document()
for s in doc.sections:
    s.top_margin = s.bottom_margin = Inches(0.7)
    s.left_margin = s.right_margin = Inches(0.8)
doc.styles["Normal"].font.name = "Calibri"
doc.styles["Normal"].font.size = Pt(10.5)


def para(text="", size=10.5, bold=False, italic=False, color=None, after=4, align=None, indent=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(after)
    if align:
        p.alignment = align
    if indent:
        p.paragraph_format.left_indent = Inches(indent)
    if text:
        r = p.add_run(text)
        r.font.size, r.bold, r.italic = Pt(size), bold, italic
        if color:
            r.font.color.rgb = color
    return p


def h1(text, accent=NAVY):
    doc.add_page_break()
    para(text, size=19, bold=True, color=accent, after=3)


def h2(text, tier=None, why=None, accent=NAVY):
    para(text, size=13.5, bold=True, color=accent, after=2)
    if tier:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(4)
        col, lab = {"***": (RED, "MUST KNOW"), "**": (AMBER, "IMPORTANT"),
                    "*": (GREY, "SEEN SOMETIMES")}[tier]
        r = p.add_run(f"{tier} {lab}  ")
        r.bold, r.font.size, r.font.color.rgb = True, Pt(9), col
        if why:
            r2 = p.add_run(why)
            r2.italic, r2.font.size, r2.font.color.rgb = True, Pt(9), GREY


def box(label, text, color):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.12)
    p.paragraph_format.space_after = Pt(5)
    r = p.add_run(f"{label}  ")
    r.bold, r.font.size, r.font.color.rgb = True, Pt(9), color
    r2 = p.add_run(text)
    r2.font.size, r2.font.color.rgb = Pt(10), color


def bullets(items, color=None):
    for it in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(2)
        r = p.add_run(it)
        r.font.size = Pt(10)
        if color:
            r.font.color.rgb = color


def table(head, rows):
    tb = doc.add_table(rows=1, cols=len(head))
    tb.style = "Light Grid Accent 1"
    tb.alignment = WD_TABLE_ALIGNMENT.CENTER
    for c, h in enumerate(head):
        tb.rows[0].cells[c].paragraphs[0].text = ""
        r = tb.rows[0].cells[c].paragraphs[0].add_run(h)
        r.bold, r.font.size = True, Pt(9.5)
    for row in rows:
        cells = tb.add_row().cells
        for c, v in enumerate(row):
            cells[c].paragraphs[0].text = ""
            cells[c].paragraphs[0].add_run(v).font.size = Pt(9.5)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)


def selftest(title, qs, color):
    para(title, size=11, bold=True, color=color, after=3)
    for i, q in enumerate(qs, 1):
        para(f"{i}. {q}", size=10, indent=0.12, after=2)


# ============================================================ COVER
para("Social Science — Study Guide", size=24, bold=True, color=NAVY, align=WD_ALIGN_PARAGRAPH.CENTER, after=2)
para("Grade 8 · for the coming test", size=13, color=NAVY, align=WD_ALIGN_PARAGRAPH.CENTER, after=2)
para("Natural Resources · Reshaping India's Political Map · Universal Franchise & Elections",
     size=11, italic=True, color=GREY, align=WD_ALIGN_PARAGRAPH.CENTER, after=10)

para("What to study first", size=12, bold=True, color=NAVY)
para("The stars are counted from the 51 real Sri Kumaran Social papers — but only where the topic is "
     "an old, recurring one. The History chapter is NEW this year, so the past papers cannot rank it; "
     "there it is marked by how much the textbook emphasises each part.", size=9.5, italic=True, color=GREY, after=6)
table(["Chapter / topic", "In the papers", "Priority"], [
    ["Geography — natural resources & conservation", "32 of 51", "*** MUST KNOW"],
    ["Geography — minerals & distribution", "18 of 51", "** IMPORTANT"],
    ["Civics — Lok Sabha / Rajya Sabha / elections", "17 of 51", "** IMPORTANT"],
    ["Civics — franchise, Election Commission, EVM", "civics core", "** IMPORTANT"],
    ["History — Delhi Sultanate, Vijayanagara, Mughals", "NEW syllabus (0 in old papers)", "** by emphasis"],
])
box("HOW TO USE THIS", "Green KEY FACT = what you must be able to state. Amber TRAP = where marks are "
    "lost. Teal FINE PRINT = the side-box detail ('Tapestry of the Past', 'Governance and Democracy') "
    "the school likes to ask. Each chapter ends with a rapid-revision list and a self-test.", NAVY)

# ============================================================ GEOGRAPHY
h1("1.  Geography — Natural Resources and their Uses", accent=GREEN)
para("The highest-scoring Social topic — natural resources appeared in 32 of the 51 past papers.",
     size=10, italic=True, color=GREY, after=6)

h2("1.1  What natural resources are, and the two ways to group them", "***",
   "Natural resources: 32 of 51 papers.", GREEN)
para("A NATURAL RESOURCE is any material or substance that occurs in Nature and is useful to humans. "
     "Trace anything around you — even a plastic button — back far enough, and it leads to Nature.")
box("KEY FACT", "Grouping 1 — by USE:  (a) resources ESSENTIAL FOR LIFE (air, water, soil/food); "
    "(b) resources for MATERIALS (wood, marble, metals); (c) resources for ENERGY (coal, water, "
    "petroleum, sunlight, wind).", GREEN)
box("KEY FACT", "Grouping 2 — RENEWABLE vs NON-RENEWABLE. RENEWABLE resources restore and regenerate "
    "in a natural rhythm (sunlight, wind, water, forests) — IF we don't disturb that rhythm by "
    "overusing them. NON-RENEWABLE resources form over very long periods and run out (coal, petroleum, "
    "minerals).", GREEN)
box("TRAP", "Renewable does NOT mean unlimited. A renewable resource can still be exhausted if we use "
    "it faster than it regenerates (e.g. groundwater).", AMBER)

h2("1.2  How resources are spread across India", "**", "Distribution/minerals: 18 of 51 papers.", GREEN)
para("Resources are spread UNEVENLY across India because of geography — iron ore, coal, oil (e.g. "
     "Mumbai High), and water lie in different regions. Sharing resources across political boundaries "
     "can cause tension and conflict.")
box("KEY FACT", "The PARADOX OF PLENTY (the 'resource curse'): a region that is RICH in resources can "
    "grow more SLOWLY if it over-relies on extracting them. Example: Punjab's groundwater is being "
    "depleted by over-pumping.", GREEN)

h2("1.3  Using resources wisely — stewardship", "***", "The biggest question set in this chapter.", GREEN)
box("KEY FACT", "STEWARDSHIP = using resources responsibly so future generations still have them. "
    "Overuse and industrial waste DISTURB Nature's cycle of restoration and regeneration.", GREEN)
bullets([
    "Organic farming — Sikkim is India's first fully organic state.",
    "Solar energy — e.g. the solar farm near Raichur, Karnataka.",
    "Rainwater harvesting; mulching to retain soil moisture; regulating fishing during the spawning season.",
])
box("FINE PRINT", "In many indigenous Indian traditions, Nature is considered SACRED — a nurturer and "
    "nourisher (e.g. Tulasī puja for wellbeing). This is a 'DON'T MISS OUT' box.", TEAL)

para("Rapid revision — Natural Resources", size=11, bold=True, color=GREEN, after=2)
bullets([
    "Natural resource = occurs in Nature + useful to humans.",
    "Group by USE: essential for life / for materials / for energy.",
    "Group by supply: renewable (regenerates) vs non-renewable (runs out).",
    "Uneven distribution → tension across boundaries; paradox of plenty = resource-rich but slow growth.",
    "Stewardship = responsible use; organic farming (Sikkim), solar (Raichur), rainwater harvesting.",
])
selftest("Self-test — Natural Resources  (answers at the end)", [
    "Name the three categories of natural resources based on their use.",
    "Give two renewable and two non-renewable resources.",
    "What is the 'paradox of plenty'?",
    "Which Indian state is fully organic, and which resource is Punjab depleting by over-pumping?",
], GREEN)

# ============================================================ HISTORY
h1("2.  History — Reshaping India's Political Map", accent=BROWN)
box("READ THIS FIRST", "This History chapter is NEW to the 2026-27 syllabus, so the old past papers do "
    "not cover it — do not be fooled by 'low frequency'. The single most useful thing you can do is "
    "learn the TIMELINE below cold; every question hangs off these dates.", RED)

h2("2.1  The timeline — learn these six dates", "**", "New syllabus; marked by emphasis.", BROWN)
table(["Year", "Event"], [
    ["1206", "The Delhi Sultanate is established"],
    ["1526", "First Battle of Panipat — Babur's victory establishes the Mughal Empire"],
    ["1556", "Second Battle of Panipat — Akbar's decisive victory"],
    ["1565", "Battle of Talikota — destruction of Vijayanagara city"],
    ["1671", "Battle of Saraighat — the Ahoms defeat the Mughals"],
    ["1699", "Formation of the Khalsa by Guru Gobind Singh"],
])

h2("2.2  The Delhi Sultanate", "**", None, BROWN)
box("KEY FACT", "The Delhi Sultanate (from 1206) was ruled by a series of dynasties, ending with the "
    "LODIS. Ala-ud-din Khilji expanded southward and repelled MONGOL invasions. Muhammad bin Tughlaq "
    "moved the capital from Delhi to DAULATABAD — an ambitious but poorly executed scheme. Timur's "
    "invasion left Delhi in ruins.", GREEN)
box("TRAP", "The JIZYA was a tax on NON-Muslim subjects in return for protection and exemption from "
    "military service. (Akbar later ABOLISHED it — see the Mughals.)", AMBER)

h2("2.3  The Vijayanagara Empire", "**", None, BROWN)
box("KEY FACT", "Founded by HARIHARA and BUKKA, capital at HAMPI. Its peak was under KRISHNADEVA RAYA "
    "— an age of military power and cultural renaissance. Portuguese travellers like Domingo Paes left "
    "detailed records. It declined after defeat at the Battle of TALIKOTA (1565).", GREEN)

h2("2.4  The Mughals", "**", None, BROWN)
box("KEY FACT", "BABUR defeated Ibrahim Lodi at the First Battle of PANIPAT (1526), founding the Mughal "
    "Empire; he wrote the Baburnama and called India a 'country of few charms'. AKBAR won the Second "
    "Battle of Panipat (1556), built a strong administration (court historian ABUL FAZL; the MIR "
    "BAKHSHI handled military matters), and ABOLISHED the jizya.", GREEN)
box("TRAP", "Sher Shah Suri briefly interrupted Mughal rule — he set up the SUR EMPIRE after defeating "
    "Humayun. And the 'towers of skulls' the chapter mentions were erected by BABUR.", AMBER)

h2("2.5  Resistance to the empires", "**", None, BROWN)
table(["Who resisted", "Key fact"], [
    ["RAJPUTS", "Maharana Pratap fought Akbar at the Battle of HALDIGHATI (1576); Rana Kumbha earlier held Mewar."],
    ["AHOMS", "Defeated the Mughals at the Battle of SARAIGHAT (1671)."],
    ["SIKHS", "Guru Gobind Singh formed the KHALSA (1699). Guru Tegh Bahadur had been executed by Aurangzeb (Gurudwara Sis Ganj Sahib marks the site)."],
])
box("FINE PRINT", "'Tapestry of the Past' details schools ask: the MUSUNURI NAYAKAS (Telugu chieftains) "
    "formed a confederacy against the Delhi Sultanate; a HUNDI was used to give payment instructions in "
    "trade networks.", TEAL)

para("Rapid revision — the Political Map", size=11, bold=True, color=BROWN, after=2)
bullets([
    "1206 Sultanate · 1526 Panipat I (Babur) · 1556 Panipat II (Akbar) · 1565 Talikota · 1671 Saraighat · 1699 Khalsa.",
    "Delhi Sultanate: Khilji (south + Mongols), Tughlaq (Daulatabad), ends with Lodis; jizya tax.",
    "Vijayanagara: Harihara & Bukka → Hampi → peak under Krishnadeva Raya → fell at Talikota.",
    "Mughals: Babur (Panipat I, Baburnama), Akbar (Panipat II, Abul Fazl, abolished jizya).",
    "Resistance: Rajputs/Haldighati, Ahoms/Saraighat, Sikhs/Khalsa.",
])
selftest("Self-test — the Political Map", [
    "Who won the First Battle of Panipat in 1526, and what did it establish?",
    "Which empire fell after the Battle of Talikota (1565), and who founded it?",
    "Who moved the capital to Daulatabad, and was it a success?",
    "Match: Haldighati, Saraighat, Khalsa → Rajputs, Ahoms, Sikhs.",
    "Which Mughal emperor abolished the jizya?",
], BROWN)

# ============================================================ CIVICS
h1("3.  Civics — Universal Franchise & India's Electoral System", accent=NAVY)
para("Elections (Lok Sabha / Rajya Sabha) appeared in 17 of the 51 past papers.",
     size=10, italic=True, color=GREY, after=6)

h2("3.1  Universal adult franchise", "**", "The foundation of the chapter.", NAVY)
box("KEY FACT", "UNIVERSAL ADULT FRANCHISE = every adult citizen (18 years and above) gets ONE vote, "
    "regardless of caste, religion, gender, wealth or education. It is guaranteed by ARTICLE 326 of "
    "the Constitution and is the cornerstone of Indian democracy.", GREEN)
para("The system bridges barriers so everyone can vote — postal/home voting for the elderly and "
     "persons with disabilities, ramps and help at booths.")

h2("3.2  The Election Commission of India", "**", None, NAVY)
box("KEY FACT", "The ELECTION COMMISSION OF INDIA (ECI) conducts elections to the Lok Sabha, State "
    "Assemblies, and the offices of President and Vice-President. It is headed by the CHIEF ELECTION "
    "COMMISSIONER. State Election Commissions handle panchayat and local-body polls.", GREEN)
table(["Officer", "Role"], [
    ["Chief Election Commissioner", "Heads the ECI; oversees the whole electoral process"],
    ["Chief Electoral Officer", "Oversees elections in a State / UT"],
    ["District Election Officer", "Oversees elections in a district"],
    ["Returning Officer (RO)", "Conducts the election in a constituency"],
    ["Electoral Registration Officer (ERO)", "Prepares and maintains the voters' list"],
])
box("FINE PRINT", "T.N. SESHAN became Chief Election Commissioner in 1990 and introduced major reforms "
    "— including voter ID cards to stop proxy voting. ('Governance and Democracy' box.)", TEAL)

h2("3.3  EVM, VVPAT and the Model Code", "**", None, NAVY)
box("KEY FACT", "Votes are cast on an EVM (Electronic Voting Machine). VVPAT = Voter Verifiable Paper "
    "Audit Trail — it prints a slip so you can confirm your vote, giving a paper trail. NOTA = 'None "
    "Of The Above', letting a voter reject all candidates.", GREEN)
box("KEY FACT", "The MODEL CODE OF CONDUCT (MCC) is the set of rules all parties must follow once "
    "elections are announced, for free and fair play — no bribing voters, no using government "
    "resources to campaign.", GREEN)

h2("3.4  Lok Sabha, Rajya Sabha, and beyond", "**", "Elections: 17 of 51 papers.", NAVY)
table(["", "How chosen", "Key facts"], [
    ["Lok Sabha", "directly elected, First-Past-the-Post", "543 constituencies; SC & ST seats reserved"],
    ["Rajya Sabha", "INDIRECTLY, by State Assemblies (single transferable vote)", "6-year term; 12 members nominated by the President"],
    ["President / VP", "INDIRECTLY, by an electoral college of MPs and MLAs", "—"],
])
box("TRAP", "FIRST-PAST-THE-POST means the candidate with the MOST votes wins — they do NOT need more "
    "than half (a majority). Lok Sabha is DIRECT; Rajya Sabha is INDIRECT — don't swap them.", AMBER)
box("FINE PRINT", "The 10th-century UTTARAMERUR inscriptions describe choosing representatives by "
    "DRAWING LOTS from a pot — an early form of elections in India. India has about 3.1 million elected "
    "representatives in local government bodies.", TEAL)

para("Rapid revision — Elections", size=11, bold=True, color=NAVY, after=2)
bullets([
    "Universal adult franchise: 1 adult (18+) = 1 vote; Article 326; cornerstone of democracy.",
    "ECI runs national/state/President polls; headed by the Chief Election Commissioner; T.N. Seshan reformed it.",
    "EVM + VVPAT (paper trail) + NOTA; Model Code of Conduct = fair-play rules once polls are announced.",
    "Lok Sabha: direct, First-Past-the-Post, 543 seats. Rajya Sabha: indirect, 6-yr term, 12 nominated.",
])
selftest("Self-test — Elections", [
    "What does universal adult franchise mean, and which Article guarantees it?",
    "Which body conducts national elections, and who heads it?",
    "What does VVPAT do, and what is NOTA?",
    "How is the Rajya Sabha elected, and how long is a member's term?",
    "Under First-Past-the-Post, does the winner need a majority?",
], NAVY)

# ============================================================ ANSWERS
h1("Self-test answers", accent=GREEN)
groups = [
    ("Natural Resources", [
        "Resources essential for life; resources for materials; resources for energy.",
        "Renewable: solar, wind (or water/forests). Non-renewable: coal, petroleum (or natural gas/minerals).",
        "A region rich in resources can grow more slowly if it over-relies on extracting them.",
        "Sikkim is fully organic; Punjab is depleting its groundwater.",
    ]),
    ("The Political Map", [
        "Babur; it established the Mughal Empire.",
        "The Vijayanagara Empire; founded by Harihara and Bukka.",
        "Muhammad bin Tughlaq; no — it was poorly executed and failed.",
        "Haldighati → Rajputs; Saraighat → Ahoms; Khalsa → Sikhs.",
        "Akbar.",
    ]),
    ("Elections", [
        "Every adult (18+) gets one vote regardless of caste/religion/gender/wealth/education; Article 326.",
        "The Election Commission of India; headed by the Chief Election Commissioner.",
        "VVPAT prints a paper slip so you can verify your vote; NOTA lets you reject all candidates.",
        "Indirectly, by the State Legislative Assemblies; a 6-year term.",
        "No — they only need the most votes, not a majority.",
    ]),
]
for title, answers in groups:
    para(title, size=12, bold=True, color=GREEN, after=2)
    for a in answers:
        para("• " + a, size=10, color=GREEN, indent=0.1, after=2)

doc.save(OUT)
print("saved:", OUT)
print(f"  {len(doc.paragraphs)} paragraphs, {len(doc.tables)} tables")
